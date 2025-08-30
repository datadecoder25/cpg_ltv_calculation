import streamlit as st
import pandas as pd
import numpy as np
import io
import zipfile
from datetime import datetime
import math
import plotly.express as px
import plotly.graph_objects as go
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import tempfile
import os
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

# Set page config
st.set_page_config(
    page_title="LTV Calculation Dashboard",
    page_icon="📊",
    layout="wide"
)

# Initialize OpenAI client (will be initialized when API key is provided)
API_KEY = '<YOUR API KEY>'
client = OpenAI(api_key = API_KEY)

def call_chat(system_prompt, prompt, model):
    """Call OpenAI Chat API with system and user prompts"""
    global client
    if not client:
        st.error("OpenAI client not initialized. Please provide an API key.")
        return None, 0, 0
    
    try:
        # Clean prompts to ensure safe encoding using comprehensive Unicode cleaning
        clean_system_prompt = clean_unicode_text(system_prompt)
        clean_user_prompt = clean_unicode_text(prompt)
        
        # Additional safety check - try to encode as ASCII to catch any remaining issues
        try:
            clean_system_prompt.encode('ascii')
            clean_user_prompt.encode('ascii')
        except UnicodeEncodeError as unicode_error:
            # Apply even more aggressive cleaning
            clean_system_prompt = ''.join(char for char in clean_system_prompt if ord(char) < 128)
            clean_user_prompt = ''.join(char for char in clean_user_prompt if ord(char) < 128)
        
        res = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": clean_system_prompt},
                {"role": "user", "content": clean_user_prompt},
            ]
        )
        return res.choices[0].message.content
    except Exception as e:
        st.error(f"Error calling OpenAI API: {str(e)}")
        return None

def process_files(uploaded_files):
    """Process uploaded CSV files and combine them into a single dataframe"""
    if not uploaded_files:
        return None
    
    dataframes = []
    
    for uploaded_file in uploaded_files:
        try:
            # Read CSV file
            df = pd.read_csv(uploaded_file, low_memory=False)
            
            # Add source file column
            df['source_file'] = uploaded_file.name
            dataframes.append(df)
            
        except Exception as e:
            st.error(f"Error reading {uploaded_file.name}: {str(e)}")
            return None
    
    if not dataframes:
        return None
    
    # Combine all dataframes
    combined_df = pd.concat(dataframes, ignore_index=True)
    
    # Replace hyphens with underscores and convert to lowercase
    combined_df.columns = combined_df.columns.str.replace('-', '_').str.replace(' ', '_').str.lower()
    
    return combined_df

def calculate_product_raw(combined_df, currency='USD'):
    """Calculate ProductRaw dataframe"""
    # Ensure purchase_date is datetime
    combined_df['purchase_date'] = pd.to_datetime(combined_df['purchase_date'], utc=True, errors='coerce')
    
    # Remove rows where purchase_date couldn't be parsed
    combined_df = combined_df.dropna(subset=['purchase_date'])
    
    # Filter data
    filtered_df = combined_df[
        (combined_df['currency'] == currency) &
        (combined_df['item_price'] > 0) &
        (combined_df['purchase_date'].dt.strftime('%Y-%m-01') > '2022-09-01')
    ].copy()
    
    # Calculate POME date
    filtered_df['buyer_email'] = filtered_df['buyer_email'].fillna('NA')
    filtered_df['pome_date'] = filtered_df.groupby(['buyer_email', 'merchant_sku'], dropna=False)['purchase_date'].transform('min').dt.strftime('%Y-%m-%d')
    
    # Create ntb dataframe
    result_df = filtered_df[['buyer_email', 'merchant_sku', 'pome_date']].drop_duplicates().reset_index(drop=True)
    ntb = result_df.rename(columns={'buyer_email': 'user_id'})
    
    # Create raw_data
    raw_data_df = filtered_df[['buyer_email', 'amazon_order_id', 'merchant_sku', 'purchase_date', 'item_price', 'shipped_quantity']].copy()
    raw_data_df = raw_data_df.rename(columns={
        'buyer_email': 'user_id',
        'amazon_order_id': 'order_id',
        'item_price': 'sales',
        'shipped_quantity': 'quantity'
    })
    
    raw_data_df['purchase_date'] = raw_data_df['purchase_date'].dt.strftime('%Y-%m-%d')
    raw_data_df['sales'] = raw_data_df['sales'].astype(float).round(2)
    raw_data = raw_data_df.drop_duplicates()
    
    # Perform left join and calculate months
    process_data = raw_data.merge(ntb, on=['user_id', 'merchant_sku'], how='left')
    process_data['purchase_date'] = pd.to_datetime(process_data['purchase_date'])
    process_data['pome_date'] = pd.to_datetime(process_data['pome_date'])
    process_data['date_diff'] = (process_data['purchase_date'] - process_data['pome_date']).dt.days
    process_data['months'] = np.ceil(process_data['date_diff'] / 30)
    
    # Group and aggregate
    process_data_agg = process_data.groupby([
        'user_id', 'merchant_sku', 'order_id', 'purchase_date', 'pome_date'
    ], dropna=False).agg({
        'sales': 'sum',
        'quantity': 'sum',
        'date_diff': 'first',
        'months': 'first'
    }).reset_index()
    
    result = process_data_agg.groupby(['merchant_sku', 'months'], dropna=False).agg({
        'user_id': 'nunique',
        'sales': 'sum',
        'quantity': 'sum',
        'order_id': 'nunique'
    }).reset_index()
    
    ProductRaw = result.rename(columns={'user_id': 'users', 'order_id': 'orders'})
    
    return ProductRaw

def calculate_product_summary(combined_df, currency='USD'):
    """Calculate ProductSummary dataframe"""
    # Create cte_tbl equivalent
    cte_tbl = combined_df[['buyer_email', 'merchant_sku', 'amazon_order_id', 'purchase_date', 'item_price', 'shipped_quantity', 'currency']].copy()
    cte_tbl['buyer_email'] = cte_tbl['buyer_email'].fillna('NA')
    
    # Apply STRFTIME transformations
    cte_tbl['purch_date'] = pd.to_datetime(cte_tbl['purchase_date']).dt.strftime('%Y-%m-%d')
    cte_tbl['purch_month'] = pd.to_datetime(cte_tbl['purchase_date']).dt.strftime('%Y-%m-01')
    
    # Apply filters
    cte_tbl = cte_tbl[
        (cte_tbl['currency'] == currency) &
        (cte_tbl['item_price'] > 0) &
        (cte_tbl['purch_month'] > '2022-09-01')
    ].copy()
    
    # Add magnitude column
    cte_tbl['magnitude'] = cte_tbl.groupby(['buyer_email', 'merchant_sku'], dropna=False)['purchase_date'].rank(method='first', ascending=True)
    
    # Create derived columns
    df = cte_tbl.copy()
    df['buyer_email'] = df['buyer_email'].astype(str)
    
    df['orig_sales'] = np.where(df['magnitude'] == 1, df['item_price'], 0)
    df['orig_quantity'] = np.where(df['magnitude'] == 1, df['shipped_quantity'], 0)
    df['orig_buyer_email_for_count'] = np.where(df['magnitude'] == 1, df['buyer_email'], np.nan)
    
    df['repeat_sales'] = np.where(df['magnitude'] > 1, df['item_price'], 0)
    df['repeat_quantity'] = np.where(df['magnitude'] > 1, df['shipped_quantity'], 0)
    df['repeat_buyer_email_for_count'] = np.where(df['magnitude'] > 1, df['buyer_email'], np.nan)
    
    # Group by and aggregate
    group_cols = ['buyer_email', 'merchant_sku', 'amazon_order_id']
    agg_operations = {
        'orig_users': ('orig_buyer_email_for_count', 'nunique'),
        'orig_sales': ('orig_sales', 'sum'),
        'orig_quantity': ('orig_quantity', 'sum'),
        'repeat_users': ('repeat_buyer_email_for_count', 'nunique'),
        'repeat_sales': ('repeat_sales', 'sum'),
        'repeat_quantity': ('repeat_quantity', 'sum')
    }
    
    cte_tbl_raw = df.groupby(group_cols, as_index=False, dropna=False).agg(**agg_operations)
    
    # Create cte_test equivalent using UNION logic
    cte_test_part1 = cte_tbl[cte_tbl['magnitude'] == 1].groupby('merchant_sku', dropna=False).agg({
        'buyer_email': 'nunique',
        'item_price': 'sum',
        'shipped_quantity': 'sum'
    }).reset_index()
    
    cte_test_part1.columns = ['merchant_sku', 'users', 'sales', 'quantity']
    cte_test_part1['users1'] = 0
    cte_test_part1['sales1'] = 0
    cte_test_part1['quantity1'] = 0
    cte_test_part1['orders'] = cte_test_part1['users']
    cte_test_part1['orders1'] = 0
    
    # Second part: magnitude > 1
    cte_test_part2 = cte_tbl[cte_tbl['magnitude'] > 1].groupby('merchant_sku', dropna=False).agg({
        'buyer_email': 'nunique',
        'item_price': 'sum',
        'shipped_quantity': 'sum'
    }).reset_index()
    
    cte_test_part2.columns = ['merchant_sku', 'users1', 'sales1', 'quantity1']
    cte_test_part2['users'] = 0
    cte_test_part2['sales'] = 0
    cte_test_part2['quantity'] = 0
    cte_test_part2['orders'] = 0
    cte_test_part2['orders1'] = cte_test_part2['users1']
    
    # Combine both parts
    cte_test = pd.concat([cte_test_part1, cte_test_part2], ignore_index=True)
    
    # Final aggregation
    ProductSummary = cte_test.groupby('merchant_sku', dropna=False).agg({
        'users': 'sum',
        'sales': 'sum',
        'quantity': 'sum',
        'users1': 'sum',
        'sales1': 'sum',
        'quantity1': 'sum',
        'orders': 'sum',
        'orders1': 'sum'
    }).reset_index()
    
    ProductSummary['orders'] = ProductSummary['quantity']
    ProductSummary['orders1'] = ProductSummary['quantity']
    
    return ProductSummary

def calculate_raw_data(combined_df, currency='USD'):
    """Calculate RawData dataframe"""
    # Create ntb equivalent
    # Convert SQL RawData query to pandas operations

    # Step 1: Create ntb equivalent
    ntb_df = combined_df[combined_df['purchase_date'].dt.strftime('%Y-%m-01') > '2022-09-01'].copy()
    ntb_df['buyer_email'] = ntb_df['buyer_email'].fillna('NA')
    ntb_df['pome_month'] = ntb_df.groupby('buyer_email', dropna=False)['purchase_date'].transform(lambda x: x.dt.strftime('%Y-%m-01').min())
    ntb = ntb_df[['buyer_email', 'pome_month']].drop_duplicates().rename(columns={'buyer_email': 'user_id'})

    # Step 2: Create pntb equivalent  
    pntb_df = combined_df[combined_df['purchase_date'].dt.strftime('%Y-%m-01') > '2022-09-01'].copy()
    pntb_df['pntb_title'] = pntb_df.groupby('merchant_sku', dropna=False)['title'].transform('min')
    pntb_df['pntb_month'] = pntb_df.groupby('merchant_sku', dropna=False)['purchase_date'].transform(lambda x: x.dt.strftime('%Y-%m-01').min())
    pntb = pntb_df[['merchant_sku', 'pntb_title', 'pntb_month']].drop_duplicates().rename(columns={'merchant_sku': 'tracking_id'})

    # Step 3: Create output_tbl equivalent
    # Filter the main dataframe
    filtered_df = combined_df[
        (combined_df['currency'] == currency) &
        (combined_df['item_price'] > 0) &
        (combined_df['purchase_date'].dt.strftime('%Y-%m-01') > '2022-09-01')
    ].copy()

    # Create month column
    filtered_df['month'] = filtered_df['purchase_date'].dt.strftime('%Y-%m-01')
    filtered_df['buyer_email'] = filtered_df['buyer_email'].fillna('NA')

    # Left join with ntb
    output_tbl1 = filtered_df.merge(
        ntb,
        left_on='buyer_email',
        right_on='user_id',
        how='left'
    )

    # Left join with pntb
    output_tbl = output_tbl1.merge(
        pntb,
        left_on='merchant_sku', 
        right_on='tracking_id',
        how='left'
    )

    output_tbl['buyer_email'] = output_tbl['buyer_email'].fillna('NA')
    output_tbl['amazon_order_id'] = output_tbl['amazon_order_id'].fillna('NA')


    # Group by and aggregate
    RawData = output_tbl.groupby(['month', 'merchant_sku', 'pntb_title', 'pome_month', 'pntb_month'], dropna=False).agg({
        'buyer_email': 'nunique',       # users
        'shipped_quantity': 'sum',      # quantity  
        'amazon_order_id': 'nunique',   # orders
        'item_price': 'sum'             # sales
    }).reset_index()

    # Rename columns
    RawData = RawData.rename(columns={
        'buyer_email': 'users',
        'shipped_quantity': 'quantity',
        'amazon_order_id': 'orders', 
        'item_price': 'sales'
    })

    # Calculate derived columns (cast as decimal(10,2) equivalent)
    RawData['avg_per_item'] = (RawData['sales'] / RawData['quantity']).round(2)
    RawData['avg_per_order'] = (RawData['sales'] / RawData['orders']).round(2)
    RawData['avg_per_user'] = (RawData['sales'] / RawData['users']).round(2)

    # Round sales to 2 decimal places (equivalent to cast as decimal(10,2))
    RawData['sales'] = RawData['sales'].round(2)
    
    return RawData

def calculate_raw_data_wo_sku(combined_df, currency='USD'):
    # Convert SQL RawData query to pandas operations
    # Step 1: Create ntb equivalent
    ntb_df_wo_sku = combined_df[combined_df['purchase_date'].dt.strftime('%Y-%m-01') > '2022-09-01'].copy()
    ntb_df_wo_sku['buyer_email'] = ntb_df_wo_sku['buyer_email'].fillna('NA')
    ntb_df_wo_sku['pome_month'] = ntb_df_wo_sku.groupby('buyer_email', dropna=False)['purchase_date'].transform(lambda x: x.dt.strftime('%Y-%m-01').min())
    ntb_df_wo_sku = ntb_df_wo_sku[['buyer_email', 'pome_month']].drop_duplicates().rename(columns={'buyer_email': 'user_id'})

    # Step 3: Create output_tbl equivalent
    # Filter the main dataframe
    filtered_df_wo_sku = combined_df[
        (combined_df['currency'] == currency) &
        (combined_df['item_price'] > 0) &
        (combined_df['purchase_date'].dt.strftime('%Y-%m-01') > '2022-09-01')
    ].copy()

    # Create month column
    filtered_df_wo_sku['month'] = filtered_df_wo_sku['purchase_date'].dt.strftime('%Y-%m-01')
    filtered_df_wo_sku['buyer_email'] = filtered_df_wo_sku['buyer_email'].fillna('NA')

    filtered_df_wo_sku = filtered_df_wo_sku[['month', 'buyer_email']].reset_index(drop=True)

    # Left join with ntb
    output_tbl_wo_sku = filtered_df_wo_sku.merge(
        ntb_df_wo_sku,
        left_on='buyer_email',
        right_on='user_id',
        how='left'
    )

    # output_tbl_wo_sku['buyer_email'] = output_tbl_wo_sku['buyer_email'].fillna('NA').drop_duplicates()


    # Group by and aggregate
    RawDataWithoutSku = output_tbl_wo_sku.groupby(['month', 'pome_month'], dropna=False).agg({
        'buyer_email': 'nunique',       # users
    }).reset_index()

    # Rename columns
    RawDataWithoutSku = RawDataWithoutSku.rename(columns={
        'buyer_email': 'users',
    })

    return RawDataWithoutSku
    

def calculate_cohort_analysis(raw_data, selected_merchant_sku=None):
    """Calculate cohort analysis"""
    # Get unique months and filter for last 12
    pome_months = raw_data['pome_month'].unique()
    pome_months_no_nan = [x for x in pome_months if pd.notna(x)]
    all_months = sorted(pome_months_no_nan)[-12:]
    
    # Apply filtering based on merchant SKU selection
    if selected_merchant_sku and selected_merchant_sku != "All":
        raw_data['IsFiltered'] = np.where(raw_data['merchant_sku'] == selected_merchant_sku, 1, 0)
        filter_msg = f"Filtering data for merchant SKU: {selected_merchant_sku}"
    else:
        raw_data['IsFiltered'] = 1
        filter_msg = "Using all merchant SKUs (no filtering applied)"
    
    # Calculate cohort sizes
    cohort_data = []
    for selected_month in all_months:
        filtered_data = raw_data[
            (raw_data['month'] == selected_month) & 
            (raw_data['pome_month'] == selected_month) & 
            (raw_data['IsFiltered'] == 1)
        ]
        cohort_size = filtered_data['users'].sum()
        cohort_data.append({
            'POME Month': selected_month,
            'Cohort Size': cohort_size
        })
    
    cohort_df = pd.DataFrame(cohort_data)
    
    # Create metrics structure
    metrics = ['Active Customers', 'Quantities', 'Orders', 'Revenue', 'Retention Rate', 'LTV', 'Cumulative LTV', 'LTV Ratio'] #'Cumulative Revenue',
    month_columns = sorted(all_months)
    
    # Create the base structure
    table_data = []
    for idx, row in cohort_df.iterrows():
        pome_month = row['POME Month']
        cohort_size = row['Cohort Size']
        
        for metric in metrics:
            row_data = {
                'POME Month': pome_month,
                'Cohort Size': cohort_size,
                'Metric': metric
            }
            
            for month in month_columns:
                row_data[month] = 0
            
            row_data['Total'] = 0
            table_data.append(row_data)
    
    filled_table = pd.DataFrame(table_data)
    
    # Convert month columns to float
    for month_col in month_columns:
        filled_table[month_col] = filled_table[month_col].astype(float)
    filled_table['Total'] = filled_table['Total'].astype(float)
    
    # Calculate metrics
    for idx, row in filled_table.iterrows():
        pome_month = row['POME Month']
        cohort_size = row['Cohort Size']
        metric = row['Metric']
        
        if cohort_size == 0:
            continue
        
        for month_col in month_columns:
            if month_col < pome_month:
                filled_table.at[idx, month_col] = 0
            else:
                filtered_data = raw_data[
                    (raw_data['month'] == month_col) & 
                    (raw_data['pome_month'] == pome_month) & 
                    (raw_data['IsFiltered'] == 1)
                ]
                
                if metric == 'Active Customers':
                    value = filtered_data['users'].sum()
                    filled_table.at[idx, month_col] = int(round(value, 0))
                elif metric == 'Quantities':
                    value = filtered_data['quantity'].sum()
                    filled_table.at[idx, month_col] = int(round(value, 0))
                elif metric == 'Orders':
                    value = filtered_data['orders'].sum()
                    filled_table.at[idx, month_col] = int(round(value, 0))
                elif metric == 'Revenue':
                    value = filtered_data['sales'].sum()
                    filled_table.at[idx, month_col] = f"${round(value, 2):.2f}"
    
    # Handle Cumulative Revenue calculation
    # for idx, row in filled_table.iterrows():
    #     if row['Metric'] == 'Cumulative Revenue':
    #         pome_month = row['POME Month']
            
    #         revenue_row_idx = None
    #         for rev_idx, rev_row in filled_table.iterrows():
    #             if (rev_row['POME Month'] == pome_month and rev_row['Metric'] == 'Revenue'):
    #                 revenue_row_idx = rev_idx
    #                 break
            
    #         if revenue_row_idx is None:
    #             continue
                
    #         for month_col in month_columns:
    #             if month_col < pome_month:
    #                 filled_table.at[idx, month_col] = 0
    #             else:
    #                 current_month_index = month_columns.index(month_col)
    #                 revenue_values = []
    #                 for future_month in month_columns[current_month_index:]:
    #                     revenue_values.append(filled_table.at[revenue_row_idx, future_month])
                    
    #                 cumulative_sum = sum(revenue_values)
    #                 filled_table.at[idx, month_col] = int(round(cumulative_sum, 0))
    
    # Handle Retention Rate calculation
    for idx, row in filled_table.iterrows():
        if row['Metric'] == 'Retention Rate':
            pome_month = row['POME Month']
            cohort_size = row['Cohort Size']
            
            for month_col in month_columns:
                if month_col < pome_month:
                    filled_table.at[idx, month_col] = 0
                elif month_col == pome_month:
                    filtered_data = raw_data[
                        (raw_data['month'] == month_col) & 
                        (raw_data['pome_month'] == pome_month) & 
                        (raw_data['IsFiltered'] == 1)
                    ]
                    
                    orders_sum = filtered_data['orders'].sum()
                    if cohort_size > 0:
                        retention_rate = (orders_sum / cohort_size) - 1
                        filled_table.at[idx, month_col] = f"{round(retention_rate*100, 2):.2f}%"
                    else:
                        filled_table.at[idx, month_col] = "0.00%"
                else:
                    active_customers_row_idx = None
                    for ac_idx, ac_row in filled_table.iterrows():
                        if (ac_row['POME Month'] == pome_month and ac_row['Metric'] == 'Active Customers'):
                            active_customers_row_idx = ac_idx
                            break
                    
                    if active_customers_row_idx is not None and cohort_size > 0:
                        active_customers = filled_table.at[active_customers_row_idx, month_col]
                        retention_rate = active_customers / cohort_size
                        filled_table.at[idx, month_col] = f"{round(retention_rate*100, 2):.2f}%"
                    else:
                        filled_table.at[idx, month_col] = "0.00%"
    
    # Handle LTV calculation
    for idx, row in filled_table.iterrows():
        if row['Metric'] == 'LTV':
            pome_month = row['POME Month']
            cohort_size = row['Cohort Size']
            
            for month_col in month_columns:
                if month_col < pome_month:
                    filled_table.at[idx, month_col] = 0
                else:
                    revenue_row_idx = None
                    for rev_idx, rev_row in filled_table.iterrows():
                        if (rev_row['POME Month'] == pome_month and rev_row['Metric'] == 'Revenue'):
                            revenue_row_idx = rev_idx
                            break
                    
                    if revenue_row_idx is not None and cohort_size > 0:
                        revenue_str = filled_table.at[revenue_row_idx, month_col]
                        # Extract numeric value from formatted revenue string
                        if isinstance(revenue_str, str) and revenue_str.startswith('$'):
                            revenue = float(revenue_str.replace('$', '').replace(',', ''))
                        else:
                            revenue = float(revenue_str) if revenue_str else 0
                        
                        ltv = revenue / cohort_size
                        filled_table.at[idx, month_col] = f"${round(ltv, 2):.2f}"
                    else:
                        filled_table.at[idx, month_col] = "$0.00"
    
    # Handle Cumulative LTV calculation
    for idx, row in filled_table.iterrows():
        if row['Metric'] == 'Cumulative LTV':
            pome_month = row['POME Month']
            cohort_size = row['Cohort Size']
            
            for month_col in month_columns:
                if month_col < pome_month:
                    filled_table.at[idx, month_col] = "$0.00"
                else:
                    # Find the corresponding LTV row for this POME Month
                    ltv_row_idx = None
                    for ltv_idx, ltv_row in filled_table.iterrows():
                        if (ltv_row['POME Month'] == pome_month and ltv_row['Metric'] == 'LTV'):
                            ltv_row_idx = ltv_idx
                            break
                    
                    if ltv_row_idx is not None:
                        # Calculate cumulative LTV from first valid month to current month
                        cumulative_ltv = 0
                        for cum_month_col in month_columns:
                            if cum_month_col < pome_month:
                                continue
                            if cum_month_col > month_col:
                                break
                                
                            ltv_str = filled_table.at[ltv_row_idx, cum_month_col]
                            if isinstance(ltv_str, str) and ltv_str.startswith('$'):
                                ltv_value = float(ltv_str.replace('$', '').replace(',', ''))
                            else:
                                ltv_value = float(ltv_str) if ltv_str else 0
                            cumulative_ltv += ltv_value
                        
                        filled_table.at[idx, month_col] = f"${round(cumulative_ltv, 2):.2f}"
                    else:
                        filled_table.at[idx, month_col] = "$0.00"
    
    # Handle LTV Ratio calculation 
    for idx, row in filled_table.iterrows():
        if row['Metric'] == 'LTV Ratio':
            pome_month = row['POME Month']
            
            # Find the corresponding Cumulative LTV row for this POME Month
            cumulative_ltv_row_idx = None
            for cum_idx, cum_row in filled_table.iterrows():
                if (cum_row['POME Month'] == pome_month and cum_row['Metric'] == 'Cumulative LTV'):
                    cumulative_ltv_row_idx = cum_idx
                    break
            
            # Find the corresponding LTV row for this POME Month to get first month LTV
            ltv_row_idx = None
            for ltv_idx, ltv_row in filled_table.iterrows():
                if (ltv_row['POME Month'] == pome_month and ltv_row['Metric'] == 'LTV'):
                    ltv_row_idx = ltv_idx
                    break
            
            if cumulative_ltv_row_idx is not None and ltv_row_idx is not None:
                # Get first month LTV (the LTV value for the POME month itself)
                first_ltv_str = filled_table.at[ltv_row_idx, pome_month]
                if isinstance(first_ltv_str, str) and first_ltv_str.startswith('$'):
                    first_ltv = float(first_ltv_str.replace('$', '').replace(',', ''))
                else:
                    first_ltv = float(first_ltv_str) if first_ltv_str else 0
                
                for month_col in month_columns:
                    if month_col < pome_month:
                        filled_table.at[idx, month_col] = "0.00x"
                    else:
                        cumulative_ltv_str = filled_table.at[cumulative_ltv_row_idx, month_col]
                        if isinstance(cumulative_ltv_str, str) and cumulative_ltv_str.startswith('$'):
                            cumulative_ltv = float(cumulative_ltv_str.replace('$', '').replace(',', ''))
                        else:
                            cumulative_ltv = float(cumulative_ltv_str) if cumulative_ltv_str else 0
                        
                        if first_ltv > 0:
                            ratio = cumulative_ltv / first_ltv
                            filled_table.at[idx, month_col] = f"{round(ratio, 2):.2f}x"
                        else:
                            filled_table.at[idx, month_col] = "0.00x"
            else:
                # If we can't find the required rows, fill with 0.00x
                for month_col in month_columns:
                    filled_table.at[idx, month_col] = "0.00x"
    
    # Calculate Total column
    for idx, row in filled_table.iterrows():
        metric = row['Metric']
        
        if metric in ['Revenue', 'LTV', 'Cumulative LTV']:
            # Handle formatted monetary values
            total_value = 0
            for month in month_columns:
                value_str = row[month]
                if isinstance(value_str, str) and value_str.startswith('$'):
                    value = float(value_str.replace('$', '').replace(',', ''))
                elif isinstance(value_str, (int, float)):
                    value = float(value_str)
                else:
                    value = 0
                total_value += value
            filled_table.at[idx, 'Total'] = f"${round(total_value, 2):.2f}"
        elif metric == 'Retention Rate':
            # Handle percentage values - calculate average retention rate
            total_values = []
            for month in month_columns:
                value_str = row[month]
                if isinstance(value_str, str) and value_str.endswith('%'):
                    value = float(value_str.replace('%', ''))
                    total_values.append(value)
                elif isinstance(value_str, (int, float)):
                    total_values.append(float(value_str))
            
            if total_values:
                avg_retention = sum(total_values) / len(total_values)
                filled_table.at[idx, 'Total'] = f"{round(avg_retention, 2):.2f}%"
            else:
                filled_table.at[idx, 'Total'] = "0.00%"
        elif metric == 'LTV Ratio':
            # Handle ratio values - calculate average ratio
            total_values = []
            for month in month_columns:
                value_str = row[month]
                if isinstance(value_str, str) and value_str.endswith('x'):
                    value = float(value_str.replace('x', ''))
                    total_values.append(value)
                elif isinstance(value_str, (int, float)):
                    total_values.append(float(value_str))
            
            if total_values:
                avg_ratio = sum(total_values) / len(total_values)
                filled_table.at[idx, 'Total'] = f"{round(avg_ratio, 2):.2f}x"
            else:
                filled_table.at[idx, 'Total'] = "0.00x"
        else:
            # Handle count metrics (Active Customers, Quantities, Orders)
            total_value = sum([row[month] for month in month_columns])
            filled_table.at[idx, 'Total'] = total_value
    
    return filled_table, filter_msg

def retention_calculation(RawData, filled_table):

    df_raw_data = RawData.copy()
    df_raw_data['month'] = pd.to_datetime(df_raw_data['month'])
    df_raw_data['pmonth_date'] = pd.to_datetime(df_raw_data['pome_month'])

    # Create a dummy df_mom_retention based on your screenshot
    mom_retention_data = {
        'Month': (list(filled_table['POME Month'])),
        'Cohort': (list(filled_table['Cohort Size']))
    }

    df_mom_retention = pd.DataFrame(mom_retention_data).drop_duplicates().reset_index(drop=True)
    df_mom_retention['Month'] = pd.to_datetime(df_mom_retention['Month'])

    def calculate_retention(cohort_month, retention_month, df_raw_data):
        """
        Calculates the MoM retention value based on the Excel formula:
        =IF(D$3<$B14,"",SUMIFS(RawData!$I:$I,RawData!$D:$D,$B14,RawData!$A:$A,D$3)/SUMIFS(RawData!$F:$F,RawData!$D:$D,$B14,RawData!$A:$A,D$3))
        
        This calculates Average Revenue Per User (ARPU) = sales/users

        Args:
            cohort_month (pd.Timestamp): The initial acquisition month of the cohort ($B14).
            retention_month (pd.Timestamp): The month for which retention is being calculated (D$3).
            raw_data_df (pd.DataFrame): The DataFrame containing the raw sales and users data.

        Returns:
            float or str: The retention value (sales/users) or an empty string if
                        retention_month is before cohort_month.
                        Returns 0 if denominator (users) is 0 to avoid division by zero.
        """

        # IF(D$3<$B14,"",...)
        if retention_month < cohort_month:
            return ""

        # Filter raw data for the specific cohort month and retention month
        # SUMIFS conditions: RawData!$D:$D=$B4 AND RawData!$A:$A=D$3
        filtered_data = df_raw_data[
            (df_raw_data['pmonth_date'] == cohort_month) &  # $D:$D=$B4 (pome_month)
            (df_raw_data['month'] == retention_month)       # $A:$A=D$3 (month)
        ]

        # SUMIFS(RawData!$I:$I,...) = sum of sales column
        sum_sales = filtered_data['sales'].sum()
        
        # SUMIFS(RawData!$F:$F,...) = sum of users column  
        sum_users = filtered_data['users'].sum()

        if sum_users == 0:
            return 0  # To handle #DIV/0! equivalent, return 0 
        else:
            return sum_sales / sum_users  # This is ARPU (Average Revenue Per User)

    # --- 3. Apply the function to populate the MoM Retention table ---

    # Get unique cohort months from your df_mom_retention
    cohort_months = sorted(df_mom_retention['Month'].unique())

    # Instead of using all months, use only the cohort analysis months (starting from 2024-07-01)
    # This matches the cohort table structure
    retention_columns_str = sorted(cohort_months)  # Use the same months as cohort analysis
    retention_columns = [pd.to_datetime(month) for month in retention_columns_str]  # Convert to datetime for comparison


    # print(f"Cohort months: {[str(d) for d in cohort_months]}")
    # print(f"Retention columns: {retention_columns_str}")

    retention_table = pd.DataFrame(index=cohort_months, columns=retention_columns)
    # retention_table
    for cohort_m in cohort_months:
        for retention_m in retention_columns_str:
            retention_table.loc[cohort_m, retention_m] = calculate_retention(
                cohort_m, retention_m, df_raw_data
            )

    retention_table = retention_table.fillna('')
    return retention_table

def retention_calculation_v2(raw_data, cohort_table):
    """
    Calculate cumulative LTV retention table from cohort analysis
    
    Args:
        raw_data (pd.DataFrame): RawData DataFrame (not used but kept for consistency)
        cohort_table (pd.DataFrame): Cohort analysis table with LTV data
        
    Returns:
        pd.DataFrame: Retention table with cumulative LTV values
    """
    
    # Extract LTV data from cohort table
    ltv_data = cohort_table[cohort_table['Metric'] == 'LTV'].copy()
    
    if ltv_data.empty:
        # If no LTV data, return empty table
        return pd.DataFrame()
    
    # Get cohort information
    cohort_info = cohort_table[cohort_table['Metric'] == 'LTV'][['POME Month', 'Cohort Size']].drop_duplicates()
    
    # Get all month columns (excluding POME Month, Cohort Size, Metric, Total)
    month_columns = [col for col in ltv_data.columns if col not in ['POME Month', 'Cohort Size', 'Metric', 'Total']]
    month_columns = sorted(month_columns)  # Ensure chronological order
    
    print(f"Creating cumulative LTV retention table...")
    print(f"Found {len(ltv_data)} cohorts and {len(month_columns)} months")
    print(f"Months: {month_columns}")
    
    # Debug: Print sample LTV data to understand format
    if not ltv_data.empty:
        print(f"Sample LTV row data:")
        sample_row = ltv_data.iloc[0]
        for col in month_columns[:3]:  # Show first 3 months
            if col in sample_row:
                print(f"  {col}: {sample_row[col]} (type: {type(sample_row[col])})")
    
    # Create the retention matrix
    cohort_months = sorted(ltv_data['POME Month'].unique())
    retention_data = []
    
    for cohort_month in cohort_months:
        # Get LTV data for this cohort
        cohort_ltv = ltv_data[ltv_data['POME Month'] == cohort_month].iloc[0]
        
        # Create row data
        row_data = {'POME Month': cohort_month}
        
        # Add cumulative LTV values for each month
        cumulative_ltv = 0
        for month_col in month_columns:
            # Only include months >= cohort month for cumulative calculation
            if month_col >= cohort_month:
                month_ltv_raw = cohort_ltv[month_col] if month_col in cohort_ltv else 0
                
                # Handle formatted LTV strings (e.g., "$12.34") and convert to numeric
                month_ltv = 0
                if isinstance(month_ltv_raw, str) and month_ltv_raw.startswith('$'):
                    try:
                        month_ltv = float(month_ltv_raw.replace('$', '').replace(',', ''))
                    except ValueError:
                        month_ltv = 0
                elif isinstance(month_ltv_raw, (int, float)):
                    month_ltv = float(month_ltv_raw)
                
                if month_ltv > 0:
                    cumulative_ltv += month_ltv
                    row_data[month_col] = f"${round(cumulative_ltv, 2):.2f}"
                else:
                    row_data[month_col] = f"${round(cumulative_ltv, 2):.2f}" if cumulative_ltv > 0 else ""
            else:
                # Months before cohort month should be empty
                row_data[month_col] = ""
        
        retention_data.append(row_data)
    
    # Create DataFrame
    retention_table = pd.DataFrame(retention_data)
    
    # Add cohort sizes
    retention_table = retention_table.merge(
        cohort_info[['POME Month', 'Cohort Size']], 
        on='POME Month', 
        how='left'
    )
    
    # Reorder columns: POME Month, Cohort Size, then month columns
    column_order = ['POME Month', 'Cohort Size'] + month_columns
    retention_table = retention_table[column_order]
    
    # Fill any remaining NaN values
    retention_table = retention_table.fillna('')
    
    print(f"✅ Cumulative LTV retention table created with shape: {retention_table.shape}")
    
    return retention_table

def create_product_ltv_table(product_raw, raw_data):
    """Create Product LTV table from merchant_sku data similar to the image format"""
    
    # Get unique merchant SKUs and create base structure
    merchant_sku_df = pd.DataFrame({"merchant_sku": list(product_raw['merchant_sku'].unique())})
    
    # Calculate users for each merchant SKU where months == 0 (first month)
    merchant_sku_df['users'] = merchant_sku_df.apply(
        lambda x: product_raw[(product_raw['months'] == 0) & (product_raw['merchant_sku'] == x['merchant_sku'])]['users'].sum(), 
        axis=1
    )
    
    # Rank by users (use 'first' method to handle ties and ensure unique ranks)
    merchant_sku_df['rank_users'] = merchant_sku_df['users'].rank(ascending=False, method='first', na_option='keep')
    
    # Set cohort equal to users
    merchant_sku_df['cohort'] = merchant_sku_df['users']
    
    # Reset index to ensure clean indexing
    merchant_sku_df = merchant_sku_df.reset_index(drop=True)
    
    # Get product titles from RawData
    def get_product_title(sku):
        """Get product title for a given SKU"""
        try:
            matching_rows = raw_data[raw_data['merchant_sku'] == sku]
            if not matching_rows.empty:
                return matching_rows['pntb_title'].iloc[0]
            else:
                return "Unknown Product"
        except:
            return "Unknown Product"
    
    # Apply the lookup safely
    merchant_sku_df['pntb_title'] = merchant_sku_df['merchant_sku'].apply(get_product_title)
    
    # Get top products by cohort size (top 10 for analysis)
    top_products = merchant_sku_df.sort_values('users', ascending=False).head(10)
    
    final_table = []
    
    for idx, (_, product_row) in enumerate(top_products.iterrows(), 1):
        sku = product_row['merchant_sku']
        cohort_size = product_row['users']
        product_title = product_row['pntb_title']
        
        # Get ProductRaw data for this SKU
        sku_data = product_raw[product_raw['merchant_sku'] == sku].copy()
        
        # Create the main product entry with all metrics
        product_entry = {
            '#': idx,
            'Product Title': product_title,
            'Cohort Size': cohort_size,
            'Merchant SKU': sku
        }
        
        # Add month columns (Month 1 through Month 12)
        for month_num in range(1, 13):
            month_idx = month_num - 1  # Convert to 0-based for data lookup
            month_data = sku_data[sku_data['months'] == month_idx]
            
            # Initialize all metrics for this month
            active_customers = month_data['users'].sum() if not month_data.empty else 0
            purchases = month_data['orders'].sum() if not month_data.empty else 0  
            revenue = int(month_data['sales'].sum()) if not month_data.empty else 0
            
            # Calculate cumulative revenue
            cumulative_revenue = 0
            for cum_month in range(0, month_idx + 1):
                cum_data = sku_data[sku_data['months'] == cum_month]
                cumulative_revenue += cum_data['sales'].sum() if not cum_data.empty else 0
            cumulative_revenue = int(cumulative_revenue)
            
            # Calculate retention rate
            retention_rate = round((active_customers / cohort_size * 100), 2) if cohort_size > 0 else 0.00
            
            # Calculate LTV
            ltv = round(revenue / cohort_size, 2) if cohort_size > 0 else 0.00
            
            # Store all metrics for this month in a structured way
            product_entry[f'Month {month_num}'] = {
                'Active Customers': active_customers,
                'Purchases': purchases,
                'Revenue': revenue,
                'Cumulative Revenue': cumulative_revenue,
                'Retention Rate': retention_rate,
                'LTV': ltv
            }
        
        final_table.append(product_entry)
    
    return final_table

def export_product_ltv_table(product_ltv_data):
    """Export Product LTV table in the exact format shown in the image"""
    
    # Create the structured table for CSV export
    csv_rows = []
    
    for product in product_ltv_data:
        # Add each metric as a separate row
        metrics = ['Active Customers', 'Purchases', 'Revenue', 'Cumulative Revenue', 'Retention Rate', 'LTV']
        
        for metric in metrics:
            metric_row = {
                '#': product['#'] if metric == 'Active Customers' else '',  # Only show # on first metric row
                'Product Title': product['Product Title'] if metric == 'Active Customers' else '',  # Only show title on first metric row  
                'Cohort Size': product['Cohort Size'] if metric == 'Active Customers' else '',  # Only show cohort size on first metric row
                'Merchant SKU': product['Merchant SKU'] if metric == 'Active Customers' else '',  # Only show SKU on first metric row
                'Metric': metric
            }
            
            # Add values for each month
            for month_num in range(1, 13):
                value = product[f'Month {month_num}'][metric]
                if metric == 'Retention Rate':
                    metric_row[f'Month {month_num}'] = f"{value}%"
                elif metric == 'LTV':
                    metric_row[f'Month {month_num}'] = f"${value}"
                else:
                    metric_row[f'Month {month_num}'] = value
            
            csv_rows.append(metric_row)
    
    # Convert to DataFrame
    export_df = pd.DataFrame(csv_rows)
    
    return export_df

def calculate_top_products_tables(product_raw, raw_data):
    """Calculate the 4 top products tables: Acquired Customers, Repeat Rate, AOV, LTV"""
    
    # Get unique merchant SKUs and create base structure
    merchant_sku_df = pd.DataFrame({"merchant_sku": list(product_raw['merchant_sku'].unique())})
    
    # Get product titles from RawData
    def get_product_title(sku):
        try:
            matching_rows = raw_data[raw_data['merchant_sku'] == sku]
            if not matching_rows.empty:
                return matching_rows['pntb_title'].iloc[0]
            else:
                return sku  # Use SKU if no title found
        except:
            return sku
    
    merchant_sku_df['product_title'] = merchant_sku_df['merchant_sku'].apply(get_product_title)
    
    # 1. Top 10 Products by Acquired Customers (Month 0 users)
    acquired_customers = []
    for _, row in merchant_sku_df.iterrows():
        sku = row['merchant_sku']
        title = row['product_title']
        
        # Get first month (month 0) users
        first_month_data = product_raw[(product_raw['merchant_sku'] == sku) & (product_raw['months'] == 0)]
        acquired = first_month_data['users'].sum() if not first_month_data.empty else 0
        
        acquired_customers.append({
            'Product Title': title,
            'Merchant SKU': sku,
            'Acquired Customers': acquired
        })
    
    top_acquired = sorted(acquired_customers, key=lambda x: x['Acquired Customers'], reverse=True)[:10]
    
    # 2. Top 10 Products by Repeat Rate
    repeat_rates = []
    for _, row in merchant_sku_df.iterrows():
        sku = row['merchant_sku']
        title = row['product_title']
        
        # Calculate repeat rate: (total customers - first month customers) / first month customers
        sku_data = product_raw[product_raw['merchant_sku'] == sku]
        first_month_users = sku_data[sku_data['months'] == 0]['users'].sum()
        total_users = sku_data['users'].sum()
        
        if first_month_users > 0:
            repeat_customers = total_users - first_month_users
            repeat_rate = (repeat_customers / first_month_users) * 100
        else:
            repeat_rate = 0
        
        repeat_rates.append({
            'Product Title': title,
            'Merchant SKU': sku,
            'Repeat Rate': f"{repeat_rate:.2f}%"
        })
    
    top_repeat = sorted(repeat_rates, key=lambda x: float(x['Repeat Rate'].replace('%', '')), reverse=True)[:10]
    
    # 3. Top 10 Products by AOV (Average Order Value)
    aovs = []
    for _, row in merchant_sku_df.iterrows():
        sku = row['merchant_sku']
        title = row['product_title']
        
        # Calculate AOV: total sales / total orders
        sku_data = product_raw[product_raw['merchant_sku'] == sku]
        total_sales = sku_data['sales'].sum()
        total_orders = sku_data['orders'].sum()
        
        aov = total_sales / total_orders if total_orders > 0 else 0
        
        aovs.append({
            'Product Title': title,
            'Merchant SKU': sku,
            'AOV': f"${aov:.2f}"
        })
    
    top_aov = sorted(aovs, key=lambda x: float(x['AOV'].replace('$', '')), reverse=True)[:10]
    
    # 4. Top 10 Products by LTV (Latest month LTV)
    ltvs = []
    for _, row in merchant_sku_df.iterrows():
        sku = row['merchant_sku']
        title = row['product_title']
        
        # Calculate LTV: total sales / first month users
        sku_data = product_raw[product_raw['merchant_sku'] == sku]
        total_sales = sku_data['sales'].sum()
        first_month_users = sku_data[sku_data['months'] == 0]['users'].sum()
        
        ltv = total_sales / first_month_users if first_month_users > 0 else 0
        
        ltvs.append({
            'Product Title': title,
            'Merchant SKU': sku,
            'LTV': f"${ltv:.2f}"
        })
    
    top_ltv = sorted(ltvs, key=lambda x: float(x['LTV'].replace('$', '')), reverse=True)[:10]
    
    return {
        'top_acquired': top_acquired,
        'top_repeat': top_repeat,
        'top_aov': top_aov,
        'top_ltv': top_ltv
    }

def clean_api_key(api_key):
    """Extract and clean OpenAI API key from chat messages or formatted text"""
    if not api_key:
        return ""
    
    import re
    
    # First, apply general Unicode cleaning
    cleaned = clean_unicode_text(api_key)
    
    # Extract OpenAI API keys using regex patterns
    # Look for patterns like sk-proj-... or sk-...
    api_key_patterns = [
        r'sk-proj-[A-Za-z0-9_-]+',  # Project API keys
        r'sk-[A-Za-z0-9_-]+',       # Regular API keys
    ]
    
    for pattern in api_key_patterns:
        matches = re.findall(pattern, cleaned)
        if matches:
            # Use the last/longest match found
            extracted_key = max(matches, key=len)
            return extracted_key
    
    # If no pattern matches, fall back to cleaned text
    return cleaned

def clean_unicode_text(text):
    """Comprehensive Unicode text cleaning for safe API transmission"""
    if not text:
        return ""
    
    import re
    
    # Convert to string if not already
    original_text = str(text)
    text = original_text
    
    # Remove or replace problematic Unicode characters
    # Narrow no-break space and related characters
    text = re.sub(r'[\u00a0\u1680\u2000-\u200f\u2028-\u202f\u205f-\u206f\u3000\ufeff]', ' ', text)
    
    # Remove other invisible/control characters
    text = re.sub(r'[\u0000-\u001f\u007f-\u009f]', '', text)
    
    # Replace smart quotes and dashes with ASCII equivalents
    text = text.replace('\u2018', "'")  # Left single quotation mark
    text = text.replace('\u2019', "'")  # Right single quotation mark
    text = text.replace('\u201c', '"')  # Left double quotation mark
    text = text.replace('\u201d', '"')  # Right double quotation mark
    text = text.replace('\u2013', '-')  # En dash
    text = text.replace('\u2014', '-')  # Em dash
    text = text.replace('\u2026', '...')  # Horizontal ellipsis
    
    # Remove any remaining non-ASCII characters
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    
    # Final safety check - brute force ASCII-only approach
    try:
        text.encode('ascii')
    except UnicodeEncodeError:
        # If still problematic, use brute force approach
        text = ''.join(char for char in text if ord(char) < 128)
    
    # Clean up excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def safe_dataframe_to_string(df, index=True):
    """Safely convert DataFrame to string, removing problematic Unicode characters"""
    if df is None or df.empty:
        return "No data available"
    
    # Convert DataFrame to string
    df_string = df.to_string(index=index)
    
    # Apply comprehensive Unicode cleaning
    return clean_unicode_text(df_string)

def extract_ltv_metrics_from_cohort(cohort_data):
    """Extract specific LTV metrics from cohort data for AI analysis"""
    if cohort_data is None or cohort_data.empty:
        return {
            'ltv_12_month': 'N/A',
            'ltv_6_month': 'N/A',
            'ltv_3_month': 'N/A',
            'first_aov_12_m': 'N/A',
            'first_aov_6_m': 'N/A',
            'first_aov_3_m': 'N/A',
            'ltv_to_aov_multiplier_12_m': 'N/A',
            'ltv_to_aov_multiplier_6_m': 'N/A',
            'ltv_to_aov_multiplier_3_m': 'N/A',
            'cohort_size_12_month': 'N/A',
            'cohort_size_6_month': 'N/A',
            'cohort_size_3_month': 'N/A',
        }
    
    try:
        print(f"🔍 METRICS DEBUG: ===== EXTRACTING LTV METRICS =====")
        print(f"🔍 METRICS DEBUG: Cohort data shape: {cohort_data.shape}")
        print(f"🔍 METRICS DEBUG: Cohort data columns: {list(cohort_data.columns)}")
        print(f"🔍 METRICS DEBUG: Available metrics: {cohort_data['Metric'].unique().tolist() if 'Metric' in cohort_data.columns else 'No Metric column'}")
        
        # Get available cohorts
        unique_cohorts = sorted(cohort_data['POME Month'].unique()) if 'POME Month' in cohort_data.columns else []
        print(f"🔍 METRICS DEBUG: Available cohorts: {unique_cohorts}")
        print(f"🔍 METRICS DEBUG: Total cohorts available: {len(unique_cohorts)}")
        
        if not unique_cohorts:
            print(f"🔍 METRICS DEBUG: ERROR - No POME Month data found")
            return {'error': 'No cohort data available'}
        
        # Select cohorts for analysis (ensuring we have enough)
        if len(unique_cohorts) < 9:
            print(f"🔍 METRICS DEBUG: WARNING - Not enough cohorts for full analysis (need at least 9, have {len(unique_cohorts)})")
            # Adjust indices based on available cohorts
            twelve_cohort_month = unique_cohorts[0]  # First cohort
            six_month_cohort_month = unique_cohorts[min(5, len(unique_cohorts)-1)]   # 6th or last available
            three_month_cohort_month = unique_cohorts[min(8, len(unique_cohorts)-1)] # 9th or last available
        else:
            twelve_cohort_month = unique_cohorts[0]  # First cohort (12 months of data)
            six_month_cohort_month = unique_cohorts[5]   # 6th cohort (6 months of data)
            three_month_cohort_month = unique_cohorts[8] # 9th cohort (3 months of data)
        
        print(f"🔍 METRICS DEBUG: Selected cohorts:")
        print(f"🔍 METRICS DEBUG: - 12-month cohort: {twelve_cohort_month}")
        print(f"🔍 METRICS DEBUG: - 6-month cohort: {six_month_cohort_month}")
        print(f"🔍 METRICS DEBUG: - 3-month cohort: {three_month_cohort_month}")
        
        # Get month columns for reference
        month_columns = [col for col in cohort_data.columns if col not in ['POME Month', 'Cohort Size', 'Metric', 'Total']]
        month_columns = [col for col in month_columns if pd.notna(col)]
        month_columns = sorted(month_columns)
        print(f"🔍 METRICS DEBUG: Available month columns: {month_columns}")
        print(f"🔍 METRICS DEBUG: Latest month column: {month_columns[-1] if month_columns else 'None'}")
        
        # Get LTV rows for each cohort
        ltv_row_12_month = cohort_data[(cohort_data['POME Month'] == twelve_cohort_month) & (cohort_data['Metric'] == 'LTV')]
        ltv_row_6_month = cohort_data[(cohort_data['POME Month'] == six_month_cohort_month) & (cohort_data['Metric'] == 'LTV')]
        ltv_row_3_month = cohort_data[(cohort_data['POME Month'] == three_month_cohort_month) & (cohort_data['Metric'] == 'LTV')]
        
        print(f"🔍 METRICS DEBUG: LTV row shapes - 12m: {ltv_row_12_month.shape}, 6m: {ltv_row_6_month.shape}, 3m: {ltv_row_3_month.shape}")
        
        if ltv_row_12_month.empty:
            print(f"🔍 METRICS DEBUG: ERROR - No LTV data for 12-month cohort {twelve_cohort_month}")
        if ltv_row_6_month.empty:
            print(f"🔍 METRICS DEBUG: ERROR - No LTV data for 6-month cohort {six_month_cohort_month}")
        if ltv_row_3_month.empty:
            print(f"🔍 METRICS DEBUG: ERROR - No LTV data for 3-month cohort {three_month_cohort_month}")
            
        if ltv_row_12_month.empty or ltv_row_6_month.empty or ltv_row_3_month.empty:
            return {'error': 'Missing LTV data for required cohorts'}

        # Extract cohort sizes
        cohort_size_12m = ltv_row_12_month.iloc[0]['Cohort Size']
        cohort_size_6m = ltv_row_6_month.iloc[0]['Cohort Size']
        cohort_size_3m = ltv_row_3_month.iloc[0]['Cohort Size']
        
        print(f"🔍 METRICS DEBUG: Cohort sizes - 12m: {cohort_size_12m}, 6m: {cohort_size_6m}, 3m: {cohort_size_3m}")

        # Extract LTV values from Total column
        ltv_12_month = ltv_row_12_month.iloc[0]['Total']
        ltv_6_month = ltv_row_6_month.iloc[0]['Total']
        ltv_3_month = ltv_row_3_month.iloc[0]['Total']
        
        print(f"🔍 METRICS DEBUG: LTV values from Total column - 12m: {ltv_12_month}, 6m: {ltv_6_month}, 3m: {ltv_3_month}")

        # Extract first AOV (LTV value where POME month = column name)
        try:
            aov_12_month = ltv_row_12_month.iloc[0][twelve_cohort_month]
            print(f"🔍 METRICS DEBUG: ✅ First AOV 12m (column {twelve_cohort_month}): {aov_12_month}")
        except KeyError:
            print(f"🔍 METRICS DEBUG: ❌ Column {twelve_cohort_month} not found for 12m AOV")
            aov_12_month = 'N/A'
            
        try:
            aov_6_month = ltv_row_6_month.iloc[0][six_month_cohort_month]
            print(f"🔍 METRICS DEBUG: ✅ First AOV 6m (column {six_month_cohort_month}): {aov_6_month}")
        except KeyError:
            print(f"🔍 METRICS DEBUG: ❌ Column {six_month_cohort_month} not found for 6m AOV")
            aov_6_month = 'N/A'
            
        try:
            aov_3_month = ltv_row_3_month.iloc[0][three_month_cohort_month]
            print(f"🔍 METRICS DEBUG: ✅ First AOV 3m (column {three_month_cohort_month}): {aov_3_month}")
        except KeyError:
            print(f"🔍 METRICS DEBUG: ❌ Column {three_month_cohort_month} not found for 3m AOV")
            aov_3_month = 'N/A'

        # Extract LTV Ratio multipliers (from LTV Ratio row, latest month column)
        latest_month = month_columns[-1] if month_columns else None
        print(f"🔍 METRICS DEBUG: Using latest month column for LTV ratios: {latest_month}")
        
        try:
            ltv_ratio_row_12m = cohort_data[(cohort_data['POME Month'] == twelve_cohort_month) & (cohort_data['Metric'] == 'LTV Ratio')]
            ltv_ratio_row_6m = cohort_data[(cohort_data['POME Month'] == six_month_cohort_month) & (cohort_data['Metric'] == 'LTV Ratio')]
            ltv_ratio_row_3m = cohort_data[(cohort_data['POME Month'] == three_month_cohort_month) & (cohort_data['Metric'] == 'LTV Ratio')]
            
            print(f"🔍 METRICS DEBUG: LTV Ratio row shapes - 12m: {ltv_ratio_row_12m.shape}, 6m: {ltv_ratio_row_6m.shape}, 3m: {ltv_ratio_row_3m.shape}")
            
            if not ltv_ratio_row_12m.empty and latest_month:
                ltv_to_aov_multiplier_12_month = ltv_ratio_row_12m.iloc[0][latest_month]
                print(f"🔍 METRICS DEBUG: ✅ LTV Ratio 12m: {ltv_to_aov_multiplier_12_month}")
            else:
                ltv_to_aov_multiplier_12_month = 'N/A'
                print(f"🔍 METRICS DEBUG: ❌ No LTV Ratio data for 12m cohort")
                
            if not ltv_ratio_row_6m.empty and latest_month:
                ltv_to_aov_multiplier_6_month = ltv_ratio_row_6m.iloc[0][latest_month]
                print(f"🔍 METRICS DEBUG: ✅ LTV Ratio 6m: {ltv_to_aov_multiplier_6_month}")
            else:
                ltv_to_aov_multiplier_6_month = 'N/A'
                print(f"🔍 METRICS DEBUG: ❌ No LTV Ratio data for 6m cohort")
                
            if not ltv_ratio_row_3m.empty and latest_month:
                ltv_to_aov_multiplier_3_month = ltv_ratio_row_3m.iloc[0][latest_month]
                print(f"🔍 METRICS DEBUG: ✅ LTV Ratio 3m: {ltv_to_aov_multiplier_3_month}")
            else:
                ltv_to_aov_multiplier_3_month = 'N/A'
                print(f"🔍 METRICS DEBUG: ❌ No LTV Ratio data for 3m cohort")
            
        except Exception as ratio_error:
            print(f"🔍 METRICS DEBUG: ERROR extracting LTV ratios: {ratio_error}")
            ltv_to_aov_multiplier_12_month = 'N/A'
            ltv_to_aov_multiplier_6_month = 'N/A'
            ltv_to_aov_multiplier_3_month = 'N/A'
        
        metrics = {
            'ltv_12_month': ltv_12_month,
            'ltv_6_month': ltv_6_month,
            'ltv_3_month': ltv_3_month,
            'first_aov_12_m': aov_12_month,
            'first_aov_6_m': aov_6_month,
            'first_aov_3_m': aov_3_month,
            'ltv_to_aov_multiplier_12_m': ltv_to_aov_multiplier_12_month,
            'ltv_to_aov_multiplier_6_m': ltv_to_aov_multiplier_6_month,
            'ltv_to_aov_multiplier_3_m': ltv_to_aov_multiplier_3_month,
            'cohort_size_12_month': cohort_size_12m,
            'cohort_size_6_month': cohort_size_6m,
            'cohort_size_3_month': cohort_size_3m,
        }
        
        print(f"🔍 METRICS DEBUG: ===== FINAL EXTRACTED METRICS =====")
        for key, value in metrics.items():
            print(f"🔍 METRICS DEBUG: {key}: {value}")
        print(f"🔍 METRICS DEBUG: ===== EXTRACTION COMPLETE =====")
        
        return metrics
        
    except Exception as e:
        print(f"🔍 METRICS DEBUG: ===== EXTRACTION ERROR =====")
        print(f"🔍 METRICS DEBUG: Error type: {type(e)}")
        print(f"🔍 METRICS DEBUG: Error message: {str(e)}")
        import traceback
        print(f"🔍 METRICS DEBUG: Full traceback: {traceback.format_exc()}")
        return {'error': f'Error extracting metrics: {str(e)}'}

def generate_product_ltv_analysis(product_sku, product_title, cohort_data, user_breakdown_data, model="gpt-4o-mini"):
    """Generate LLM analysis for a specific product's LTV data"""
    
    cohort_metrics = extract_ltv_metrics_from_cohort(cohort_data)
    first_order_aov_12_months = cohort_metrics['first_aov_12_m']
    twelve_month_ltv = cohort_metrics['ltv_12_month']
    twelve_month_ltv_ratio = cohort_metrics['ltv_to_aov_multiplier_12_m']
    first_order_aov_6_months = cohort_metrics['first_aov_6_m']
    six_month_ltv = cohort_metrics['ltv_6_month']
    six_month_ltv_ratio = cohort_metrics['ltv_to_aov_multiplier_6_m']
    first_order_aov_3_months = cohort_metrics['first_aov_3_m']  
    three_month_ltv = cohort_metrics['ltv_3_month']
    three_month_ltv_ratio = cohort_metrics['ltv_to_aov_multiplier_3_m']
    cohort_size_12_month = cohort_metrics['cohort_size_12_month']
    cohort_size_6_month = cohort_metrics['cohort_size_6_month']
    cohort_size_3_month = cohort_metrics['cohort_size_3_month']

    
    # System prompt with the theory and instructions
    system_prompt = f"""You are an experienced CMO and analytics lead for CPG brands. 
    Read the attached cohort LTV data table and produce a concise, executive-ready report. 
    Follow every instruction below precisely. 

Your task is to analyze product-level LTV data and provide insights following this framework:

What you have
•	An account-level cohort table (CSV) with a column and monthly period columns (e.g., = month of entry, then P1 … P12).
•	The Metric column includes exactly these rows (case-insensitive match): Active Customers, Quantities, Orders, Revenue, Retention Rate, LTV.
•	If any row is missing, state that clearly and proceed using what’s available.

Output structure (use these section headings)
    Title: LTV Review — Account Level (User Lifecycle Analysis)
    Executive Summary (bullets, 5–8 lines max)

    cohort_size_12m = {cohort_size_12_month}
    first_order_aov_12m = {first_order_aov_12_months}
    12 month LTV = {twelve_month_ltv}
    12 month LTV to first month AOV Ratio = {twelve_month_ltv_ratio}

    cohort_size_6m = {cohort_size_6_month}
    first_order_aov_6m = {first_order_aov_6_months}
    6 month LTV = {six_month_ltv}
    6 month LTV to first month AOV Ratio = {six_month_ltv_ratio}

    cohort_size_3m = {cohort_size_3_month}
    first_order_aov_3m = {first_order_aov_3_months}
    3 month LTV = {three_month_ltv}
    3 month LTV to first month AOV Ratio = {three_month_ltv_ratio}   

    Cohort walk-throughs (exactly how to pick them)
    Find and discuss three cohorts:
    1.	12-Month Cohort
    2.	6-Month Cohort
    3.	3-Month Cohort
    For each cohort, cover these points (short paragraph + 3–5 bullets):
    •	NTB size: total unique NTBs in P0 (the cohort size).
    •	Repeat behavior: how many of them repeated over the window and the % of the original cohort; call out any strong months or drop-offs.
    •	Commercial impact: resulting Orders, Units, and Revenue (summarize at the horizon).
    •	LTV vs first order: LTV multiple at the horizon (e.g., “12-mo LTV is 3.8× first-order AOV”).
    Keep the narration instructional: show how to read the table and why each metric matters.

    Benchmark check (use this JSON)
    Compare the 12-month account-level LTV multiple (not 6- or 3-month) to the first-order AOV against these category benchmarks. If you know the brand’s category, use it; if not, choose the closest match and say which you used.
    [
    "category":"Snacks / pantry","cadence":"3–6 orders/yr","healthy_min_x":2.5,"healthy_max_x":4,"elite_min_x":4,"elite_max_x":5,
    "category":"Supplements (caps/gummies)","cadence":"monthly–bi-monthly","healthy_min_x":4,"healthy_max_x":7,"elite_min_x":7,"elite_max_x":9,
    "category":"Sports nutrition (protein, pre-/post-)","cadence":"6–10 orders/yr","healthy_min_x":4,"healthy_max_x":8,"elite_min_x":8,"elite_max_x":10,
    "category":"Coffee / tea / pods","cadence":"monthly–bi-monthly","healthy_min_x":4,"healthy_max_x":7,"elite_min_x":7,"elite_max_x":10,
    "category":"Hydration / greens powders","cadence":"monthly","healthy_min_x":4,"healthy_max_x":7,"elite_min_x":7,"elite_max_x":10,
    "category":"Beverages, RTD (shelf-stable)","cadence":"3–6 orders/yr","healthy_min_x":3,"healthy_max_x":5,"elite_min_x":5,"elite_max_x":7,
    "category":"Beauty & personal care (consumable)","cadence":"3–6 orders/yr","healthy_min_x":3,"healthy_max_x":6,"elite_min_x":6,"elite_max_x":9,
    "category":"Oral care (paste, brush heads)","cadence":"3–6 orders/yr","healthy_min_x":3,"healthy_max_x":5,"elite_min_x":5,"elite_max_x":8,
    "category":"Household cleaning & laundry","cadence":"3–6 orders/yr","healthy_min_x":3,"healthy_max_x":6,"elite_min_x":6,"elite_max_x":9,
    "category":"Pet consumables (food, treats, litter)","cadence":"4–8 orders/yr","healthy_min_x":4,"healthy_max_x":7,"elite_min_x":7,"elite_max_x":10,
    "category":"Baby care (diapers, wipes, formula)","cadence":"8–12 orders/yr","healthy_min_x":5,"healthy_max_x":9,"elite_min_x":9,"elite_max_x":12,
    "category":"Condiments & sauces","cadence":"2–4 orders/yr","healthy_min_x":2,"healthy_max_x":4,"elite_min_x":4,"elite_max_x":6,
    "category":"Spices / baking","cadence":"2–4 orders/yr","healthy_min_x":2,"healthy_max_x":3.5,"elite_min_x":3.5,"elite_max_x":5
    ]
    How to state it:
    “12-mo LTV multiple = 3.2× vs first-order AOV. For [Chosen Category], that sits in the healthy range (2.5–4.0×) and below elite (≥4.0×).”


    Implications & Next Moves
    End with 5–8 bullet recommendations tailored to what you see, e.g.:
    •	If 12-mo LTV multiple is healthy/elite, note room to raise CAC (or maintain) to capture share.
    •	If weak, propose product-level cohort dive; shift budget toward SKUs with higher LTV.
    •	If repeat rate is front-loaded, test post-purchase CX, subscriptions, and review pipelines.
    •	If P0 AOV is low, test bundles/upsells to lift first-order economics.
    •	If retention decays after a specific month, propose a lifecycle nudge (email/remarketing/loyalty).
    •	Recommend AMC audiences based on high-propensity repeaters.

    Guardrails
    •	Do not hallucinate. If any row/column is missing or ambiguous, state it and explain the adjustment.
    •	State assumptions (e.g., whether LTV in the table is “per customer” or “total”; use the table labels to decide).
    •	Prefer short, skimmable paragraphs and bullets.
    •	Show only the essential numbers for decisions (NTB size, repeat %, LTV multiple, and revenue at the horizon).
    •	Use the brand’s currency as displayed.

    Data handling notes (if needed)
    •	First-Order AOV = P0 Revenue / P0 Orders for the cohort.
    •	If LTV is not provided but cumulative Revenue is, compute Cumulative Revenue ÷ Cohort size at P0.
    •	Repeaters = Count of unique customers who purchased in months P1+ (from Active Customers), expressed as % of P0.
    •	When selecting the 12-, 6-, 3-month cohorts, confirm the last available month in the table and choose cohorts with at least that many periods of follow-up.

"""

    # Create user prompt with the actual data
    cleaned_title = clean_unicode_text(product_title)
    cleaned_sku = clean_unicode_text(product_sku)
    cohort_string = safe_dataframe_to_string(cohort_data)
    
    user_prompt = f"""Analyze the LTV performance for this product:

PRODUCT: {cleaned_title}
SKU: {cleaned_sku}

COHORT ANALYSIS DATA:
{cohort_string}

Focus on actionable insights with clear formatting, bullet points, and professional structure suitable for executive presentations."""

    # Call the LLM
    analysis = call_chat(clean_unicode_text(system_prompt), user_prompt, model)
    
    return {
        'analysis': analysis
    }

def generate_top_products_table_explanation(top_products_data, model="gpt-4o-mini"):
    """Generate LLM explanation for the top 10 products tables"""
    

    
    system_prompt = """You are an experienced e-commerce analyst. Analyze the top 10 product performance tables and provide clear, actionable insights.

Your task is to explain what these four tables reveal about product performance and provide strategic recommendations.

Output structure:
## Product Portfolio Performance Analysis

### Key Insights from Top Product Rankings
- Analyze patterns across the four ranking categories
- Identify products that appear in multiple top 10 lists
- Note any surprises or standout performers

### Strategic Observations
- Comment on the relationship between acquired customers, repeat rate, AOV, and LTV
- Identify which products drive volume vs. value
- Point out any misalignments (high acquisition but low LTV, etc.)

### Portfolio Optimization Recommendations
- Suggest which products deserve more marketing investment
- Identify products that may need attention (low repeat rates, low LTV despite high acquisition)
- Recommend cross-sell or bundling opportunities

Keep the analysis concise, executive-ready, and focused on actionable insights. Use bullet points for readability."""

    # Format the data for the prompt
    acquired_table = pd.DataFrame(top_products_data['top_acquired'])
    repeat_table = pd.DataFrame(top_products_data['top_repeat'])
    aov_table = pd.DataFrame(top_products_data['top_aov'])
    ltv_table = pd.DataFrame(top_products_data['top_ltv'])

    user_prompt = f"""Analyze these four top 10 product performance tables:

TOP 10 PRODUCTS BY ACQUIRED CUSTOMERS:
{safe_dataframe_to_string(acquired_table, index=False)}

TOP 10 PRODUCTS BY REPEAT RATE:
{safe_dataframe_to_string(repeat_table, index=False)}

TOP 10 PRODUCTS BY AOV:
{safe_dataframe_to_string(aov_table, index=False)}

TOP 10 PRODUCTS BY LTV:
{safe_dataframe_to_string(ltv_table, index=False)}

Provide strategic insights and recommendations based on these rankings."""

    analysis = call_chat(clean_unicode_text(system_prompt), user_prompt, model)
    return {'analysis': analysis}

def generate_user_breakdown_explanation(user_breakdown_data, model="gpt-4o-mini"):
    """Generate LLM explanation for user breakdown analysis"""
    

    
    system_prompt = """You are an experienced customer analytics specialist. Analyze the new vs old user breakdown data and provide clear insights about customer acquisition and retention patterns.

Your task is to explain what the user breakdown data reveals about customer behavior and business health.

Output structure:
## User Acquisition & Retention Analysis

### Customer Base Composition
- Summarize the overall split between new and old users
- Identify monthly trends in new user acquisition
- Comment on old user retention patterns

### Business Health Indicators
- Assess whether the business is growing (more new users) or mature (more repeat users)
- Identify any concerning trends in user acquisition or retention
- Comment on seasonal patterns if visible

### Strategic Implications
- Recommend actions based on new vs old user ratios
- Suggest focus areas for marketing and retention efforts
- Identify months that may need special attention

Keep the analysis concise, data-driven, and focused on actionable business insights."""

    user_prompt = f"""Analyze this user breakdown data showing new vs old users over time:

USER BREAKDOWN DATA:
{safe_dataframe_to_string(user_breakdown_data, index=False)}

Provide insights about customer acquisition patterns, retention trends, and business health indicators based on this data."""

    analysis = call_chat(clean_unicode_text(system_prompt), user_prompt, model)
    return {'analysis': analysis}

def generate_account_ltv_analysis(cohort_data, user_breakdown_data, model="gpt-4o-mini"):
    """Generate LLM analysis for account-level LTV data"""
    


    try:
        cohort_metrics = extract_ltv_metrics_from_cohort(cohort_data)
        first_order_aov_12_months = cohort_metrics['first_aov_12_m']
        twelve_month_ltv = cohort_metrics['ltv_12_month']
        twelve_month_ltv_ratio = cohort_metrics['ltv_to_aov_multiplier_12_m']
        first_order_aov_6_months = cohort_metrics['first_aov_6_m']
        six_month_ltv = cohort_metrics['ltv_6_month']
        six_month_ltv_ratio = cohort_metrics['ltv_to_aov_multiplier_6_m']
        first_order_aov_3_months = cohort_metrics['first_aov_3_m']  
        three_month_ltv = cohort_metrics['ltv_3_month']
        three_month_ltv_ratio = cohort_metrics['ltv_to_aov_multiplier_3_m']
        cohort_size_12_month = cohort_metrics['cohort_size_12_month']
        cohort_size_6_month = cohort_metrics['cohort_size_6_month']
        cohort_size_3_month = cohort_metrics['cohort_size_3_month']
    except Exception as e:
        print("Error in generate_account_ltv_analysis:", e)
    # System prompt with the theory and instructions
    system_prompt = f"""You are an experienced CMO and analytics lead for CPG brands. 
    Read the attached cohort LTV data table and produce a concise, executive-ready report. 
    Follow every instruction below precisely. 

Your task is to analyze account-level LTV data and provide insights following this framework:

Output structure (use these section headings)
    Title: LTV Review — Account Level (User Lifecycle Analysis)
    Executive Summary (bullets, 5–8 lines max, write it based on the "Cohort Walk-through", "Benchmark check" and "Implications & Next moves" analysis that you will do below.)
    Cohort Walk-throughs
    •	A) 12-Month Cohort (oldest with ≥12 months of follow-up)
    •	B) 6-Month Cohort (month with ≥6 months of follow-up)
    •	C) 3-Month Cohort (month with ≥3 months of follow-up)
    Benchmark Check (12-Month LTV vs First-Order AOV)
    Implications & Next Moves (bullets)
    Keep the tone crisp and businesslike. Use only the most decision-relevant numbers. Format key callouts as bold* and multiples as ****e.g., 3.2×****. Round money to whole currency (no cents), percentages to 1 decimal, and multiples to 1 decimal unless precision is needed.*

    # Note: Educational content about LTV has been moved to the report as a standalone section

    first_order_aov_12_months = {first_order_aov_12_months}
    12 month LTV = {twelve_month_ltv}
    12 month LTV to first month AOV Ratio = {twelve_month_ltv_ratio}
    cohort_size_12m = {cohort_size_12_month}
    first_order_aov_6_months = {first_order_aov_6_months}
    6 month LTV = {six_month_ltv}
    6 month LTV to first month AOV Ratio = {six_month_ltv_ratio}
    cohort_size_6m = {cohort_size_6_month}
    first_order_aov_3_months = {first_order_aov_3_months}
    3 month LTV = {three_month_ltv}
    3 month LTV to first month AOV Ratio = {three_month_ltv_ratio}
    cohort_size_3m = {cohort_size_3_month}

    Cohort walk-throughs (exactly how to pick them)
    12-Month Cohort Analysis
    Thoroughly review the first POME month data set across all 6 metrics, Active Customers, Quantities, Orders, Revenue, Retention Rate, and LTV, covering a 12-month account-level window. The report should include the following:
    •	Total number of unique NTBs acquired in the P0 month, which is the cohort size
    •	Repeat behavior across months:
    Track how many of these customers returned in each subsequent month, P1, P2, P3, P6, P12. Report the repeat rate (%) month by month and identify any noticeable patterns. Did the repeat rate drop sharply? If so, after which month? Was the drop-off consistent, or did it stabilize? Based on this, assess whether the repeat purchase behavior is front-loaded.
    •	Performance contribution:
    Break down how many orders these NTBs generated, how many units they purchased, and how much revenue they contributed, starting from P0, then compare to P2, P3, P6, and P12. Explain how each metric evolved over time.
    •	LTV Analysis:
    Compare the 3-month LTV to the average first purchase, how many times greater is it? Do the same for the 6-month and 12-month LTV. Express each as an “X times” multiple relative to the first purchase. 
    You will find the info in the COHORT ANALYSIS DATA table.

    6-Month Cohort Analysis
    Repeat the same structure for the first POME month of the 6-month cohort:
    •	Total unique NTBs in the P0 month,which is the cohort size
    •	Repeat rate month by month:
    Detail how many NTBs repeated in P1, P2, P3, and P6, with their respective percentages. Comment on when the repeat rate drops sharply, if at all, and whether the pattern suggests front-loaded behavior.
    •	Performance impact:
    Highlight total orders, units sold, and revenue generated—starting from P0, then P2, P3, and P6. Show how these numbers changed over time.
    •	LTV Commentary:
    Compare the 3-month and 6-month LTV figures to the first order value. Calculate how many X the LTV is vs. the first purchase.

    3-Month Cohort Analysis
    Again, apply the same reporting logic to the 3-month cohort:
    •	Total NTBs in the P0 month, which is the cohort size
    •	Retention trends:
    Explain how many repeated in P1, P2, and P3, along with percentages. Analyze if and when the drop-offs happen and whether the repeat pattern is front-loaded.
    •	Performance outcomes:
    Narrate how many orders, units, and how much revenue was driven by these NTBs over P0, P2, and P3.
    •	LTV Evaluation:
    Compare the 3-month LTV to the first purchase value and express the uplift as a multiple.


    Compare the 12-month account-level LTV multiple (not 6- or 3-month) to the first-order AOV against these category benchmarks. If you know the brand’s category, use it; if not, choose the closest match and say which you used.
    [
    "category":"Snacks / pantry","cadence":"3–6 orders/yr","healthy_min_x":2.5,"healthy_max_x":4,"elite_min_x":4,"elite_max_x":5,
    "category":"Supplements (caps/gummies)","cadence":"monthly–bi-monthly","healthy_min_x":4,"healthy_max_x":7,"elite_min_x":7,"elite_max_x":9,
    "category":"Sports nutrition (protein, pre-/post-)","cadence":"6–10 orders/yr","healthy_min_x":4,"healthy_max_x":8,"elite_min_x":8,"elite_max_x":10,
    "category":"Coffee / tea / pods","cadence":"monthly–bi-monthly","healthy_min_x":4,"healthy_max_x":7,"elite_min_x":7,"elite_max_x":10,
    "category":"Hydration / greens powders","cadence":"monthly","healthy_min_x":4,"healthy_max_x":7,"elite_min_x":7,"elite_max_x":10,
    "category":"Beverages, RTD (shelf-stable)","cadence":"3–6 orders/yr","healthy_min_x":3,"healthy_max_x":5,"elite_min_x":5,"elite_max_x":7,
    "category":"Beauty & personal care (consumable)","cadence":"3–6 orders/yr","healthy_min_x":3,"healthy_max_x":6,"elite_min_x":6,"elite_max_x":9,
    "category":"Oral care (paste, brush heads)","cadence":"3–6 orders/yr","healthy_min_x":3,"healthy_max_x":5,"elite_min_x":5,"elite_max_x":8,
    "category":"Household cleaning & laundry","cadence":"3–6 orders/yr","healthy_min_x":3,"healthy_max_x":6,"elite_min_x":6,"elite_max_x":9,
    "category":"Pet consumables (food, treats, litter)","cadence":"4–8 orders/yr","healthy_min_x":4,"healthy_max_x":7,"elite_min_x":7,"elite_max_x":10,
    "category":"Baby care (diapers, wipes, formula)","cadence":"8–12 orders/yr","healthy_min_x":5,"healthy_max_x":9,"elite_min_x":9,"elite_max_x":12,
    "category":"Condiments & sauces","cadence":"2–4 orders/yr","healthy_min_x":2,"healthy_max_x":4,"elite_min_x":4,"elite_max_x":6,
    "category":"Spices / baking","cadence":"2–4 orders/yr","healthy_min_x":2,"healthy_max_x":3.5,"elite_min_x":3.5,"elite_max_x":5
    ]
    How to state it:
    “12-mo LTV multiple = 3.2× vs first-order AOV. For [Chosen Category], that sits in the healthy range (2.5–4.0×) and below elite (≥4.0×).”

    Implications & Next Moves
    End with 5–8 bullet recommendations tailored to what you see, e.g.:
    •	If 12-mo LTV multiple is healthy/elite, note room to raise CAC (or maintain) to capture share.
    •	If weak, propose product-level cohort dive; shift budget toward SKUs with higher LTV.
    •	If repeat rate is front-loaded, test post-purchase CX, subscriptions, and review pipelines.
    •	If P0 AOV is low, test bundles/upsells to lift first-order economics.
    •	If retention decays after a specific month, propose a lifecycle nudge (email/remarketing/loyalty).
    •	Recommend AMC audiences based on high-propensity repeaters.


    Guardrails
    •	Do not hallucinate. If any row/column is missing or ambiguous, state it and explain the adjustment.
    •	State assumptions (e.g., whether LTV in the table is “per customer” or “total”; use the table labels to decide).
    •	Prefer short, skimmable paragraphs and bullets.
    •	Show only the essential numbers for decisions (NTB size, repeat %, LTV multiple, and revenue at the horizon).
    •	Use the brand’s currency as displayed.


    Data handling notes (if needed)
    •	First-Order AOV = P0 Revenue / P0 Orders for the cohort.
    •	If LTV is not provided but cumulative Revenue is, compute Cumulative Revenue ÷ Cohort size at P0.
    •	Repeaters = Count of unique customers who purchased in months P1+ (from Active Customers), expressed as % of P0.
    •	When selecting the 12-, 6-, 3-month cohorts, confirm the last available month in the table and choose cohorts with at least that many periods of follow-up.


"""

    # Create user prompt with the actual data
    user_prompt = f"""Analyze the account-level LTV performance across all products:

SCOPE: Account-Level Analysis

COHORT ANALYSIS DATA:
{safe_dataframe_to_string(cohort_data)}

Focus on actionable insights with clear formatting, bullet points, and professional structure suitable for executive presentations. Pay special attention to the Cumulative LTV and LTV Ratio metrics to understand how customer value compounds over time."""

    # Call the LLM
    analysis = call_chat(clean_unicode_text(system_prompt), user_prompt, model)
    
    return {
        'analysis': analysis
    }

def calculate_user_breakdown(raw_data, raw_data_wo_sku, selected_merchant_sku=None):
    """
    Calculate old users vs new users breakdown based on merchant SKU selection
    
    Args:
        raw_data (pd.DataFrame): RawData DataFrame
        selected_merchant_sku (str): Selected merchant SKU or None for all
        
    Returns:
        pd.DataFrame: DataFrame with month, all_users, new_users, old_users
    """
    
    # Apply SKU filtering if specified
    if selected_merchant_sku and selected_merchant_sku != "All":
        filtered_data = raw_data[raw_data['merchant_sku'] == selected_merchant_sku].copy()
    else:
        filtered_data = raw_data_wo_sku.copy()
    
    # Calculate new users (where pome_month equals month - first-time buyers)
    new_users_data = filtered_data[filtered_data['pome_month'] == filtered_data['month']].copy()
    new_users = new_users_data.groupby(['month']).agg({'users': 'sum'}).reset_index().rename(columns={'users': 'new_users'})
    
    # Calculate all users for each month
    all_users = filtered_data.groupby(['month']).agg({'users': 'sum'}).reset_index().rename(columns={'users': 'all_users'})
    
    # Combine and calculate old users
    combined_users = pd.merge(all_users, new_users, on='month', how='left')
    combined_users['new_users'] = combined_users['new_users'].fillna(0)
    combined_users['old_users'] = combined_users['all_users'] - combined_users['new_users']
    
    # Sort by month
    combined_users['month'] = pd.to_datetime(combined_users['month'])
    combined_users = combined_users.sort_values('month').reset_index(drop=True)
    
    return combined_users

def create_user_breakdown_chart(user_breakdown_df, selected_sku):
    """
    Create a stacked bar chart showing old users vs new users
    
    Args:
        user_breakdown_df (pd.DataFrame): DataFrame with user breakdown data
        selected_sku (str): Selected merchant SKU for title
        
    Returns:
        plotly.graph_objects.Figure: Plotly figure object
    """
    

    
    # Create stacked bar chart
    fig = go.Figure()
    
    # Add new users bar
    fig.add_trace(go.Bar(
        name='New Users',
        x=user_breakdown_df['month'].dt.strftime('%Y-%m'),
        y=user_breakdown_df['new_users'],
        marker_color='#2E86AB',  # Blue
        text=user_breakdown_df['new_users'],
        textposition='inside',
        texttemplate='%{text}',
        hovertemplate='<b>New Users</b><br>' +
                      'Month: %{x}<br>' +
                      'Count: %{y:,}<extra></extra>'
    ))
    
    # Add old users bar
    fig.add_trace(go.Bar(
        name='Returning Users',
        x=user_breakdown_df['month'].dt.strftime('%Y-%m'),
        y=user_breakdown_df['old_users'],
        marker_color='#A23B72',  # Purple/Pink
        text=user_breakdown_df['old_users'],
        textposition='inside',
        texttemplate='%{text}',
        hovertemplate='<b>Returning Users</b><br>' +
                      'Month: %{x}<br>' +
                      'Count: %{y:,}<extra></extra>'
    ))
    
    # Update layout for stacked bar chart
    sku_text = f" - {selected_sku}" if selected_sku != "All" else " - All SKUs"
    
    fig.update_layout(
        title=f'User Breakdown by Month{sku_text}',
        xaxis_title='Month',
        yaxis_title='Number of Users',
        barmode='stack',
        hovermode='x unified',
        height=500,
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        ),
        margin=dict(t=80, b=40, l=40, r=40),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)'
    )
    
    # Style the axes
    fig.update_xaxes(
        showgrid=True,
        gridwidth=1,
        gridcolor='rgba(128,128,128,0.2)',
        tickangle=45
    )
    
    fig.update_yaxes(
        showgrid=True,
        gridwidth=1,
        gridcolor='rgba(128,128,128,0.2)'
    )
    
    return fig

def clean_text_for_pdf(text):
    """Clean text to make it safe for PDF generation with ReportLab"""
    if not text:
        return ""
    
    import re
    
    # Convert markdown headers to HTML headers (order matters - start with most specific)
    text = re.sub(r'^#### (.*?)$', r'<br/><b><u>\1</u></b><br/>', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.*?)$', r'<br/><b><u>\1</u></b><br/>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.*?)$', r'<br/><br/><b><font size="14">\1</font></b><br/>', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.*?)$', r'<br/><br/><b><font size="16">\1</font></b><br/><br/>', text, flags=re.MULTILINE)
    
    # Convert markdown-style bold to HTML bold (be more specific to avoid conflicts)
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    
    # Convert markdown-style italic to HTML italic (avoid conflicts with bold)
    text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'<i>\1</i>', text)
    
    # Handle markdown lists
    text = re.sub(r'^[-*+] (.*?)$', r'- \1', text, flags=re.MULTILINE)
    text = re.sub(r'^\d+\. (.*?)$', r'• \1', text, flags=re.MULTILINE)
    
    # Replace bullet points with simple dashes
    text = text.replace('•', '- ').replace('◦', '- ').replace('—', '- ')
    
    # Remove emojis and special unicode characters that might cause issues
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    
    # Clean up any remaining problematic characters
    text = text.replace('\\', '')
    text = text.replace('"', "'")
    text = text.replace('`', "'")  # Remove backticks
    
    # Handle code blocks and inline code
    text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)  # Remove code blocks
    text = re.sub(r'`([^`]+)`', r'\1', text)  # Remove inline code formatting
    
    # Clean up horizontal rules
    text = re.sub(r'^[-*_]{3,}$', '', text, flags=re.MULTILINE)
    
    # Ensure proper line breaks
    text = text.replace('\n\n', '<br/><br/>').replace('\n', '<br/>')
    
    # Remove any nested paragraph tags that might cause issues
    text = re.sub(r'<para.*?>', '', text)
    text = text.replace('</para>', '')
    
    # Clean up multiple consecutive break tags
    text = re.sub(r'(<br/>){3,}', '<br/><br/>', text)
    
    return text

def clean_text_for_word(text):
    """Clean text to make it properly formatted for Word documents"""
    if not text:
        return ""
    
    import re
    
    # Remove markdown headers and replace with plain text (order matters - most specific first)
    text = re.sub(r'^#### (.*?)$', r'\n\1\n', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.*?)$', r'\n\1\n', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.*?)$', r'\n\1\n', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.*?)$', r'\n\1\n\n', text, flags=re.MULTILINE)
    
    # Convert markdown-style bold to plain text (remove ** markers)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Convert markdown-style italic to plain text (remove * markers, avoid conflicts with bold)
    text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'\1', text)
    
    # Handle markdown lists
    text = re.sub(r'^[-*+] (.*?)$', r'- \1', text, flags=re.MULTILINE)
    text = re.sub(r'^\d+\. (.*?)$', r'- \1', text, flags=re.MULTILINE)
    
    # Remove code blocks and inline code
    text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)  # Remove code blocks
    text = re.sub(r'`([^`]+)`', r'\1', text)  # Remove inline code formatting
    
    # Clean up horizontal rules
    text = re.sub(r'^[-*_]{3,}$', '', text, flags=re.MULTILINE)
    
    # Clean up bullet points
    text = text.replace('•', '- ').replace('◦', '- ').replace('—', '- ')
    
    # Remove emojis and special unicode characters that might cause issues
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    
    # Clean up any remaining problematic characters
    text = text.replace('\\', '')
    text = text.replace('`', '')  # Remove remaining backticks
    
    # Clean up extra whitespace and line breaks
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Replace multiple line breaks with double
    text = re.sub(r'^\s+|\s+$', '', text)  # Trim whitespace from start/end
    
    return text

def generate_product_data_files(product_raw, raw_data):
    """Generate individual data files for each of the top 10 products"""
    
    # Get top 10 products
    top_products_data = calculate_top_products_tables(product_raw, raw_data)
    top_10_products = top_products_data['top_acquired'][:10]
    
    # Create a ZIP file with all product data
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for product in top_10_products:
            product_sku = product['Merchant SKU']
            product_title = product['Product Title']
            
            # Clean SKU for filename
            safe_sku = "".join(c for c in product_sku if c.isalnum() or c in (' ', '-', '_')).rstrip()
            
            try:
                # Generate cohort analysis for this product
                product_cohort_table, _ = calculate_cohort_analysis(raw_data, product_sku)
                if not product_cohort_table.empty:
                    zip_file.writestr(f"{safe_sku}_Cohort_Analysis.csv", product_cohort_table.to_csv(index=False))
                
                # Generate user breakdown for this product
                product_user_breakdown = calculate_user_breakdown(raw_data, raw_data, product_sku)
                if not product_user_breakdown.empty:
                    zip_file.writestr(f"{safe_sku}_User_Breakdown.csv", product_user_breakdown.to_csv(index=False))
                
                # Generate product LTV data for this product
                product_ltv_data = create_product_ltv_table(product_raw[product_raw['merchant_sku'] == product_sku], raw_data)
                if product_ltv_data:
                    product_ltv_export = export_product_ltv_table(product_ltv_data)
                    zip_file.writestr(f"{safe_sku}_Product_LTV.csv", product_ltv_export.to_csv(index=False))
                    
            except Exception as e:
                # Create error file if data generation fails
                error_content = f"Error generating data for {product_title} (SKU: {product_sku}): {str(e)}"
                zip_file.writestr(f"{safe_sku}_ERROR.txt", error_content)
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def generate_comprehensive_ltv_report(product_raw, raw_data, cohort_table, user_breakdown_df, model="gpt-4o-mini"):
    """Generate a comprehensive LTV report with all components"""
    
    # Create a temporary file for the comprehensive report
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    doc = SimpleDocTemplate(temp_file.name, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#2c3e50')
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=15,
        spaceBefore=25,
        textColor=colors.HexColor('#34495e')
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=15,
        textColor=colors.HexColor('#2c3e50')
    )
    
    analysis_style = ParagraphStyle(
        'AnalysisStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=12,
        spaceBefore=10,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor('#2c3e50'),
        leftIndent=20,
        rightIndent=20
    )
    
    # Title Page
    elements.append(Paragraph("Comprehensive LTV Analysis Report", title_style))
    elements.append(Spacer(1, 30))
    elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
    elements.append(Spacer(1, 30))
    
    # 1. Account-Level LTV Analysis
    elements.append(Paragraph("1. Account-Level LTV Analysis", heading_style))
    
    if client:  # Only generate if OpenAI client is available
        try:
            account_analysis = generate_account_ltv_analysis(cohort_table, user_breakdown_df, model)
            if account_analysis['analysis']:
                analysis_text = clean_text_for_pdf(account_analysis['analysis'])
                elements.append(Paragraph(analysis_text, analysis_style))
                elements.append(Spacer(1, 20))
        except Exception as e:
            elements.append(Paragraph(f"Account LTV analysis unavailable: {str(e)}", analysis_style))
            elements.append(Spacer(1, 15))
    else:
        elements.append(Paragraph("Account-level LTV analysis is available when an OpenAI API key is provided.", analysis_style))
        elements.append(Spacer(1, 20))
    
    elements.append(PageBreak())
    
    # 2. Top 10 Products Analysis
    elements.append(Paragraph("2. Top Product Performance Tables", heading_style))
    
    # Calculate top products data
    top_products_data = calculate_top_products_tables(product_raw, raw_data)
    
    # Create tables for each category
    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ])
    
    # Top 10 by Acquired Customers
    elements.append(Paragraph("2.1 Top 10 Products by Acquired Customers", subheading_style))
    acquired_data = [['Rank', 'Product Title', 'Merchant SKU', 'Acquired Customers']]
    if top_products_data.get('top_acquired'):
        for i, product in enumerate(top_products_data['top_acquired'], 1):
            acquired_data.append([str(i), product['Product Title'][:40] + '...' if len(product['Product Title']) > 40 else product['Product Title'], 
                                 product['Merchant SKU'], str(product['Acquired Customers'])])
    else:
        acquired_data.append(['N/A', 'No data available', 'N/A', '0'])
    
    acquired_table = Table(acquired_data, colWidths=[0.8*inch, 3*inch, 1.5*inch, 1.2*inch])
    acquired_table.setStyle(table_style)
    elements.append(acquired_table)
    elements.append(Spacer(1, 15))
    
    # Top 10 by Repeat Rate
    elements.append(Paragraph("2.2 Top 10 Products by Repeat Rate", subheading_style))
    repeat_data = [['Rank', 'Product Title', 'Merchant SKU', 'Repeat Rate']]
    if top_products_data.get('top_repeat'):
        for i, product in enumerate(top_products_data['top_repeat'], 1):
            repeat_data.append([str(i), product['Product Title'][:40] + '...' if len(product['Product Title']) > 40 else product['Product Title'], 
                               product['Merchant SKU'], product['Repeat Rate']])
    else:
        repeat_data.append(['N/A', 'No data available', 'N/A', '0%'])
    
    repeat_table = Table(repeat_data, colWidths=[0.8*inch, 3*inch, 1.5*inch, 1.2*inch])
    repeat_table.setStyle(table_style)
    elements.append(repeat_table)
    elements.append(Spacer(1, 15))
    
    # Top 10 by AOV
    elements.append(Paragraph("2.3 Top 10 Products by AOV", subheading_style))
    aov_data = [['Rank', 'Product Title', 'Merchant SKU', 'AOV']]
    if top_products_data.get('top_aov'):
        for i, product in enumerate(top_products_data['top_aov'], 1):
            aov_data.append([str(i), product['Product Title'][:40] + '...' if len(product['Product Title']) > 40 else product['Product Title'], 
                            product['Merchant SKU'], product['AOV']])
    else:
        aov_data.append(['N/A', 'No data available', 'N/A', '$0.00'])
    
    aov_table = Table(aov_data, colWidths=[0.8*inch, 3*inch, 1.5*inch, 1.2*inch])
    aov_table.setStyle(table_style)
    elements.append(aov_table)
    elements.append(Spacer(1, 15))
    
    # Top 10 by LTV
    elements.append(Paragraph("2.4 Top 10 Products by LTV", subheading_style))
    ltv_data = [['Rank', 'Product Title', 'Merchant SKU', 'LTV']]
    if top_products_data.get('top_ltv'):
        for i, product in enumerate(top_products_data['top_ltv'], 1):
            ltv_data.append([str(i), product['Product Title'][:40] + '...' if len(product['Product Title']) > 40 else product['Product Title'], 
                            product['Merchant SKU'], product['LTV']])
    else:
        ltv_data.append(['N/A', 'No data available', 'N/A', '$0.00'])
    
    ltv_table = Table(ltv_data, colWidths=[0.8*inch, 3*inch, 1.5*inch, 1.2*inch])
    ltv_table.setStyle(table_style)
    elements.append(ltv_table)
    elements.append(Spacer(1, 20))
    
    # Top Products Table Analysis
    elements.append(Paragraph("2.5 Product Portfolio Analysis", subheading_style))
    if client:
        try:
            tables_analysis = generate_top_products_table_explanation(top_products_data, model)
            if tables_analysis['analysis']:
                analysis_text = clean_text_for_pdf(tables_analysis['analysis'])
                elements.append(Paragraph(analysis_text, analysis_style))
                elements.append(Spacer(1, 20))
        except Exception as e:
            elements.append(Paragraph(f"Product portfolio analysis unavailable: {str(e)}", analysis_style))
            elements.append(Spacer(1, 15))
    else:
        elements.append(Paragraph("Product portfolio analysis is available when an OpenAI API key is provided.", analysis_style))
        elements.append(Spacer(1, 20))
    
    elements.append(PageBreak())
    
    # 3. Individual Product Analysis for Top 10
    elements.append(Paragraph("3. Individual Product LTV Analysis", heading_style))
    elements.append(Paragraph("Detailed analysis for each of the top 10 products by acquired customers:", styles['Normal']))
    elements.append(Spacer(1, 15))
    
    # Get top 10 products by acquired customers for individual analysis
    top_10_products = top_products_data['top_acquired'][:10] if top_products_data.get('top_acquired') else []
    
    if not top_10_products:
        elements.append(Paragraph("No products found for individual analysis.", styles['Normal']))
        elements.append(Spacer(1, 15))
    else:
        for i, product in enumerate(top_10_products, 1):
            product_sku = product['Merchant SKU']
            product_title = product['Product Title']
            
            elements.append(Paragraph(f"3.{i} {product_title}", subheading_style))
            elements.append(Paragraph(f"SKU: {product_sku}", styles['Normal']))
            elements.append(Spacer(1, 10))
            
            # Generate individual product analysis
            if client:
                try:
                    product_cohort_table, _ = calculate_cohort_analysis(raw_data, product_sku)
                    product_user_breakdown = calculate_user_breakdown(raw_data, raw_data, product_sku)
                    
                    product_analysis = generate_product_ltv_analysis(product_sku, product_title, product_cohort_table, product_user_breakdown, model)
                    if product_analysis['analysis']:
                        analysis_text = clean_text_for_pdf(product_analysis['analysis'])
                        elements.append(Paragraph(analysis_text, analysis_style))
                        elements.append(Spacer(1, 15))
                except Exception as e:
                    elements.append(Paragraph(f"Analysis unavailable for {product_title}: {str(e)}", analysis_style))
                    elements.append(Spacer(1, 10))
            else:
                elements.append(Paragraph("Individual product analysis is available when an OpenAI API key is provided.", analysis_style))
                elements.append(Spacer(1, 15))
            
            if i < len(top_10_products):  # Don't add page break after last product
                elements.append(PageBreak())
    
    # 4. User Breakdown Analysis
    elements.append(PageBreak())
    elements.append(Paragraph("4. User Acquisition & Retention Analysis", heading_style))
    
    if user_breakdown_df is not None and not user_breakdown_df.empty:
        # Add key metrics summary
        elements.append(Paragraph("4.1 Key Metrics Summary", subheading_style))
        
        total_new_users = user_breakdown_df['new_users'].sum()
        total_returning_users = user_breakdown_df['old_users'].sum()
        total_all_users = user_breakdown_df['all_users'].sum()
        # new_user_percentage = (total_new_users / total_all_users) * 100 if total_all_users > 0 else 0
        # avg_new_users_per_month = user_breakdown_df['new_users'].mean()
        # avg_returning_users_per_month = user_breakdown_df['old_users'].mean()
        
        metrics_data = [
            ['Metric', 'Value'],
            ['Total New Users', f"{total_new_users:,}"],
            ['Total Returning Users', f"{total_returning_users:,}"],
            ['Total Users', f"{total_all_users:,}"],
            # ['New User Percentage', f"{new_user_percentage:.1f}%"],
            # ['Avg New Users/Month', f"{avg_new_users_per_month:.0f}"],
            # ['Avg Returning Users/Month', f"{avg_returning_users_per_month:.0f}"]
        ]
        
        metrics_table = Table(metrics_data, colWidths=[3*inch, 2*inch])
        metrics_table.setStyle(table_style)
        elements.append(metrics_table)
        elements.append(Spacer(1, 20))
        
        # Add user breakdown chart
        elements.append(Paragraph("4.2 User Breakdown Visualization", subheading_style))
        try:
            # Create the chart
            fig = create_user_breakdown_chart(user_breakdown_df, "All")
            
            # Save chart as image and add to PDF
            import tempfile
            chart_temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            
            try:
                # Try to save as image using kaleido
                fig.write_image(chart_temp_file.name, format='png', width=800, height=500, scale=2)
                
                from reportlab.platypus import Image
                chart_image = Image(chart_temp_file.name, width=6*inch, height=3.75*inch)
                elements.append(chart_image)
                elements.append(Spacer(1, 20))
                
            except Exception as e:
                # Check if it's a kaleido-specific error
                if 'kaleido' in str(e).lower():
                    error_msg = f"Chart generation unavailable: Kaleido package error - {str(e)}"
                else:
                    error_msg = f"Chart generation unavailable: {str(e)}"
                
                elements.append(Paragraph(error_msg, analysis_style))
                elements.append(Paragraph("Please ensure kaleido is installed: pip install --upgrade kaleido", analysis_style))
                elements.append(Spacer(1, 15))
            finally:
                # Clean up temp file
                try:
                    os.unlink(chart_temp_file.name)
                except:
                    pass
                    
        except Exception as e:
            elements.append(Paragraph(f"Chart generation error: {str(e)}", analysis_style))
            elements.append(Spacer(1, 15))
        
        # Add user breakdown data table
        elements.append(Paragraph("4.3 Detailed User Breakdown Data", subheading_style))
        
        # Prepare table data
        table_data = [['Month', 'Total Users', 'New Users', 'Returning Users', 'New User %', 'Returning User %']]
        
        for _, row in user_breakdown_df.iterrows():
            month_str = row['month'].strftime('%Y-%m-%d') if hasattr(row['month'], 'strftime') else str(row['month'])
            total_users = int(row['all_users'])
            new_users = int(row['new_users'])
            returning_users = int(row['old_users'])
            new_pct = (new_users / total_users * 100) if total_users > 0 else 0
            returning_pct = (returning_users / total_users * 100) if total_users > 0 else 0
            
            table_data.append([
                month_str,
                f"{total_users:,}",
                f"{new_users:,}",
                f"{returning_users:,}",
                f"{new_pct:.1f}%",
                f"{returning_pct:.1f}%"
            ])
        
        # Create table with appropriate column widths
        breakdown_table = Table(table_data, colWidths=[1.2*inch, 1*inch, 1*inch, 1.2*inch, 0.8*inch, 1*inch])
        breakdown_table.setStyle(table_style)
        elements.append(breakdown_table)
        elements.append(Spacer(1, 20))
        
        # Analysis insights
        if len(user_breakdown_df) >= 2:
            elements.append(Paragraph("4.4 Analysis Insights", subheading_style))
            
            recent_new_users = user_breakdown_df.tail(3)['new_users'].mean()
            if len(user_breakdown_df) > 3:
                earlier_new_users = user_breakdown_df.head(len(user_breakdown_df)-3)['new_users'].mean()
            else:
                # Safety check: make sure we have at least 1 row
                if len(user_breakdown_df) >= 1:
                    earlier_new_users = user_breakdown_df.head(1)['new_users'].iloc[0]
                else:
                    earlier_new_users = 0
            
            total_months = len(user_breakdown_df)
            months_with_growth = len(user_breakdown_df[user_breakdown_df['new_users'] > user_breakdown_df['old_users']])
            if len(user_breakdown_df) > 0:
                best_month = user_breakdown_df.loc[user_breakdown_df['new_users'].idxmax()]
            else:
                best_month = {'month': 'N/A', 'new_users': 0}
            
            insights_text = f"""
            • Analysis covers {total_months} months of user acquisition data
            • {months_with_growth} out of {total_months} months showed more new users than returning users
            • Best acquisition month: {best_month['month'].strftime('%Y-%m') if hasattr(best_month['month'], 'strftime') else str(best_month['month'])} with {best_month['new_users']:.0f} new users
            """
            
            if earlier_new_users > 0:
                new_user_trend = ((recent_new_users - earlier_new_users) / earlier_new_users) * 100
                trend_direction = "increasing" if new_user_trend > 5 else "decreasing" if new_user_trend < -5 else "stable"
                insights_text += f"• New user acquisition trend: {trend_direction} ({new_user_trend:+.1f}%)"
            
            elements.append(Paragraph(insights_text, analysis_style))
            elements.append(Spacer(1, 20))
    
    # AI-generated analysis (if available)
    if client and user_breakdown_df is not None:
        try:
            elements.append(Paragraph("4.5 AI-Generated Analysis", subheading_style))
            user_analysis = generate_user_breakdown_explanation(user_breakdown_df, model)
            if user_analysis['analysis']:
                analysis_text = clean_text_for_pdf(user_analysis['analysis'])
                elements.append(Paragraph(analysis_text, analysis_style))
                elements.append(Spacer(1, 20))
        except Exception as e:
            elements.append(Paragraph(f"AI user breakdown analysis unavailable: {str(e)}", analysis_style))
            elements.append(Spacer(1, 15))
    elif user_breakdown_df is None or user_breakdown_df.empty:
        elements.append(Paragraph("User breakdown analysis is unavailable - no user breakdown data provided.", analysis_style))
        elements.append(Spacer(1, 20))
    else:
        elements.append(Paragraph("AI-generated user breakdown analysis is available when an OpenAI API key is provided.", analysis_style))
        elements.append(Spacer(1, 20))
    
    # Build PDF
    doc.build(elements)
    temp_file.close()
    
    return temp_file.name

def generate_comprehensive_word_report(product_raw, raw_data, cohort_table, user_breakdown_df, model="gpt-4o-mini"):
    """Generate a comprehensive LTV report in Word format with all components"""
    
    # Create a temporary file for the comprehensive report
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_file.close()  # Close immediately so docx can write to it
    
    # Create Word document
    doc = Document()
    
    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('Comprehensive LTV Analysis Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.name = 'Calibri'
    
    doc.add_paragraph()  # Add space
    
    # Date
    date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Educational Content - Understanding LTV
    edu_heading = doc.add_heading('Understanding and Applying LTV', level=1)
    edu_heading.runs[0].font.size = Pt(16)
    edu_heading.runs[0].font.name = 'Calibri'
    
    # LTV Introduction
    ltv_intro = doc.add_paragraph(
        "Lifetime Value (LTV) is a core metric used to assess the long-term value of customers acquired by a business. "
        "It's a foundational KPI for evaluating the effectiveness of marketing efforts, particularly in online retail, "
        "where LTV is often compared to Customer Acquisition Cost (CAC) to determine the ROI of advertising spend."
    )
    ltv_intro.style = 'Normal'
    
    ltv_definition = doc.add_paragraph(
        "At its simplest, LTV represents the total purchases a customer makes over a defined time period. "
        "It's critical that this timeline is explicitly stated, whether 6, 12, or 24 months, so LTV comparisons "
        "across channels, time periods, or businesses are accurate and meaningful."
    )
    ltv_definition.style = 'Normal'
    
    ltv_calculation = doc.add_paragraph(
        "LTV is calculated by grouping customers into cohorts and tracking their purchase behavior over time. "
        "This cohort-based approach allows for a more accurate and actionable view of how customer value evolves."
    )
    ltv_calculation.style = 'Normal'
    
    # Common Pitfall
    pitfall_heading = doc.add_heading('A Common Pitfall to Avoid', level=2)
    pitfall_heading.runs[0].font.size = Pt(14)
    pitfall_heading.runs[0].font.name = 'Calibri'
    
    pitfall_para = doc.add_paragraph(
        "One frequent mistake in calculating LTV is using an undifferentiated pool of customers over a fixed data range "
        "(e.g., 12 months), then summing purchases without considering when each customer entered the system. "
        "This approach is flawed because customers acquired in the later months haven't had enough time to make repeat purchases."
    )
    pitfall_para.style = 'Normal'
    
    pitfall_solution = doc.add_paragraph(
        "LTV must always be tied to a specific cohort, a group of customers acquired at the same time, "
        "and measured over a fixed timeline. Only then can we accurately assess repeat behavior and long-term value."
    )
    pitfall_solution.style = 'Normal'
    
    # Why This Exercise Matters
    importance_heading = doc.add_heading('Why This Exercise Matters', level=2)
    importance_heading.runs[0].font.size = Pt(14)
    importance_heading.runs[0].font.name = 'Calibri'
    
    importance_intro = doc.add_paragraph(
        "This isn't just about building a dashboard, it's about uncovering insights that drive business growth. Here's how:"
    )
    importance_intro.style = 'Normal'
    
    # Use Cases
    use_case1_heading = doc.add_paragraph("1. Improved Budgeting for Customer Acquisition")
    use_case1_heading.style = 'Normal'
    for run in use_case1_heading.runs:
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    use_case1_content = doc.add_paragraph(
        "A weight loss supplement brand had a $22 average first order value (AOV) and was acquiring customers at a $2 ROAS, "
        "equivalent to $11 per customer. Initially, this looked unsustainable, with 50% of ad revenue going to acquisition. "
        "However, once they calculated the 12-month LTV, they found it reached $105. With this insight, they realized they "
        "could confidently maintain or even increase acquisition spend to accelerate growth."
    )
    use_case1_content.style = 'Normal'
    
    use_case2_heading = doc.add_paragraph("2. Product-Level Optimization")
    use_case2_heading.style = 'Normal'
    for run in use_case2_heading.runs:
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    use_case2_content = doc.add_paragraph(
        "Analyzing LTV by product helps identify which SKUs drive higher retention and long-term value. "
        "This allows the brand to shift more ad budget toward high-LTV products, improving both profitability and sustainability."
    )
    use_case2_content.style = 'Normal'
    
    use_case3_heading = doc.add_paragraph("3. Avoiding Misguided Spend")
    use_case3_heading.style = 'Normal'
    for run in use_case3_heading.runs:
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    use_case3_content = doc.add_paragraph(
        "If a flagship product has poor LTV, say just 1.5x its initial AOV, it may not justify additional ad investment. "
        "Instead, this insight signals a need to investigate root causes (e.g., low customer satisfaction or product ratings) "
        "before increasing spend."
    )
    use_case3_content.style = 'Normal'
    
    use_case4_heading = doc.add_paragraph("4. Smarter Audience Targeting via AMC")
    use_case4_heading.style = 'Normal'
    for run in use_case4_heading.runs:
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    use_case4_content = doc.add_paragraph(
        "Using LTV-based cohorts, brands can create custom audiences within Amazon Marketing Cloud and activate them "
        "through both Sponsored and DSP campaigns, driving smarter, higher-ROI media strategies."
    )
    use_case4_content.style = 'Normal'
    
    # Cohort Definition Methodology
    methodology_heading = doc.add_heading('Cohort Definition Methodology', level=2)
    methodology_heading.runs[0].font.size = Pt(14)
    methodology_heading.runs[0].font.name = 'Calibri'
    
    methodology_content = doc.add_paragraph(
        "For account-level analysis (aggregated across all products), we define a cohort as unique new-to-brand customers "
        "acquired in a specific month, who haven't purchased from the brand in the past 12 months (assuming access to 24 months of data). "
        "Their purchasing behavior is then tracked monthly for up to 12 months to generate cohort-specific insights."
    )
    methodology_content.style = 'Normal'
    
    # How to Read This Table
    table_reading_heading = doc.add_heading('How to Read This Table', level=2)
    table_reading_heading.runs[0].font.size = Pt(14)
    table_reading_heading.runs[0].font.name = 'Calibri'
    
    # Table reading bullet points
    table_bullets = [
        "Active Customers: Number of unique NTB customers who purchased in that month.",
        "Orders / Quantities / Revenue: Sales from that cohort in that month.",
        "Retention Rate: % of the original cohort active in that month (or % of prior month - use the table's definition; state which you see).",
        "LTV: Cumulative revenue per original cohort customer through that month (if instead it's total cumulative revenue, divide by the cohort size at P0)."
    ]
    
    for bullet in table_bullets:
        bullet_para = doc.add_paragraph(f"• {bullet}")
        bullet_para.style = 'Normal'
        for run in bullet_para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
    
    doc.add_page_break()
    
    # 1. Account-Level LTV Analysis
    heading1 = doc.add_heading('1. Account-Level LTV Analysis', level=1)
    heading1.runs[0].font.size = Pt(16)
    heading1.runs[0].font.name = 'Calibri'
    
    if client:  # Only generate if OpenAI client is available
        try:
            account_analysis = generate_account_ltv_analysis(cohort_table, user_breakdown_df, model)
            if account_analysis['analysis']:
                cleaned_analysis = clean_text_for_word(account_analysis['analysis'])
                analysis_para = doc.add_paragraph(cleaned_analysis)
                analysis_para.style = 'Normal'
                for run in analysis_para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
        except Exception as e:
            error_para = doc.add_paragraph(f"Account LTV analysis unavailable: {str(e)}")
            error_para.style = 'Normal'
    else:
        placeholder_para = doc.add_paragraph("Account-level LTV analysis is available when an OpenAI API key is provided.")
        placeholder_para.style = 'Normal'
    
    doc.add_page_break()
    
    # 2. Top 10 Products Analysis
    heading2 = doc.add_heading('2. Top Product Performance Tables', level=1)
    heading2.runs[0].font.size = Pt(16)
    heading2.runs[0].font.name = 'Calibri'
    
    # Calculate top products data
    top_products_data = calculate_top_products_tables(product_raw, raw_data)
    
    # Helper function to create tables
    def create_word_table(doc, data, title):
        if title and title.strip():  # Only create heading if title is not empty
            subheading = doc.add_heading(title, level=2)
            if subheading.runs:  # Check if runs exist before accessing
                subheading.runs[0].font.size = Pt(14)
                subheading.runs[0].font.name = 'Calibri'
        
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(data[0]):
            header_cells[i].text = str(header)
            # Apply formatting if runs exist
            if header_cells[i].paragraphs[0].runs:
                header_cells[i].paragraphs[0].runs[0].bold = True
                header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        
        # Data rows
        for row_idx, row_data in enumerate(data[1:], 1):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data)
                # Apply formatting if runs exist
                if row_cells[col_idx].paragraphs[0].runs:
                    row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)
        
        doc.add_paragraph()  # Add space after table
    
    # Top 10 by Acquired Customers
    acquired_data = [['Rank', 'Product Title', 'Merchant SKU', 'Acquired Customers']]
    for i, product in enumerate(top_products_data['top_acquired'], 1):
        acquired_data.append([
            str(i), 
            product['Product Title'][:50] + '...' if len(product['Product Title']) > 50 else product['Product Title'], 
            product['Merchant SKU'], 
            str(product['Acquired Customers'])
        ])
    create_word_table(doc, acquired_data, '2.1 Top 10 Products by Acquired Customers')
    
    # Top 10 by Repeat Rate
    repeat_data = [['Rank', 'Product Title', 'Merchant SKU', 'Repeat Rate']]
    for i, product in enumerate(top_products_data['top_repeat'], 1):
        repeat_data.append([
            str(i), 
            product['Product Title'][:50] + '...' if len(product['Product Title']) > 50 else product['Product Title'], 
            product['Merchant SKU'], 
            product['Repeat Rate']
        ])
    create_word_table(doc, repeat_data, '2.2 Top 10 Products by Repeat Rate')
    
    # Top 10 by AOV
    aov_data = [['Rank', 'Product Title', 'Merchant SKU', 'AOV']]
    for i, product in enumerate(top_products_data['top_aov'], 1):
        aov_data.append([
            str(i), 
            product['Product Title'][:50] + '...' if len(product['Product Title']) > 50 else product['Product Title'], 
            product['Merchant SKU'], 
            product['AOV']
        ])
    create_word_table(doc, aov_data, '2.3 Top 10 Products by AOV')
    
    # Top 10 by LTV
    ltv_data = [['Rank', 'Product Title', 'Merchant SKU', 'LTV']]
    for i, product in enumerate(top_products_data['top_ltv'], 1):
        ltv_data.append([
            str(i), 
            product['Product Title'][:50] + '...' if len(product['Product Title']) > 50 else product['Product Title'], 
            product['Merchant SKU'], 
            product['LTV']
        ])
    create_word_table(doc, ltv_data, '2.4 Top 10 Products by LTV')
    
    # Top Products Table Analysis
    subheading_portfolio = doc.add_heading('2.5 Product Portfolio Analysis', level=2)
    subheading_portfolio.runs[0].font.size = Pt(14)
    subheading_portfolio.runs[0].font.name = 'Calibri'
    
    if client:
        try:
            tables_analysis = generate_top_products_table_explanation(top_products_data, model)
            if tables_analysis['analysis']:
                cleaned_analysis = clean_text_for_word(tables_analysis['analysis'])
                portfolio_para = doc.add_paragraph(cleaned_analysis)
                portfolio_para.style = 'Normal'
                for run in portfolio_para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
        except Exception as e:
            error_para = doc.add_paragraph(f"Product portfolio analysis unavailable: {str(e)}")
            error_para.style = 'Normal'
    else:
        placeholder_para = doc.add_paragraph("Product portfolio analysis is available when an OpenAI API key is provided.")
        placeholder_para.style = 'Normal'
    
    doc.add_page_break()
    
    # 3. Individual Product Analysis for Top 10
    heading3 = doc.add_heading('3. Individual Product LTV Analysis', level=1)
    heading3.runs[0].font.size = Pt(16)
    heading3.runs[0].font.name = 'Calibri'
    
    intro_para = doc.add_paragraph("Detailed analysis for each of the top 10 products by acquired customers:")
    intro_para.style = 'Normal'
    
    # Get top 10 products by acquired customers for individual analysis
    top_10_products = top_products_data['top_acquired'][:10] if top_products_data.get('top_acquired') else []
    
    if not top_10_products:
        no_products_para = doc.add_paragraph("No products found for individual analysis.")
        no_products_para.style = 'Normal'
    else:
        for i, product in enumerate(top_10_products, 1):
            product_sku = product['Merchant SKU']
            product_title = product['Product Title']
            
            product_heading = doc.add_heading(f'3.{i} {product_title}', level=2)
            product_heading.runs[0].font.size = Pt(14)
            product_heading.runs[0].font.name = 'Calibri'
            
            sku_para = doc.add_paragraph(f"SKU: {product_sku}")
            sku_para.style = 'Normal'
            
            # Generate individual product analysis
            if client:
                try:
                    product_cohort_table, _ = calculate_cohort_analysis(raw_data, product_sku)
                    product_user_breakdown = calculate_user_breakdown(raw_data, raw_data, product_sku)
                    
                    product_analysis = generate_product_ltv_analysis(product_sku, product_title, product_cohort_table, product_user_breakdown, model)
                    if product_analysis['analysis']:
                        cleaned_analysis = clean_text_for_word(product_analysis['analysis'])
                        analysis_para = doc.add_paragraph(cleaned_analysis)
                        analysis_para.style = 'Normal'
                        for run in analysis_para.runs:
                            run.font.size = Pt(11)
                            run.font.name = 'Calibri'
                except Exception as e:
                    error_para = doc.add_paragraph(f"Analysis unavailable for {product_title}: {str(e)}")
                    error_para.style = 'Normal'
            else:
                placeholder_para = doc.add_paragraph("Individual product analysis is available when an OpenAI API key is provided.")
                placeholder_para.style = 'Normal'
            
            if i < len(top_10_products):  # Don't add page break after last product
                doc.add_page_break()
    
    # 4. User Breakdown Analysis
    doc.add_page_break()
    heading4 = doc.add_heading('4. User Acquisition & Retention Analysis', level=1)
    heading4.runs[0].font.size = Pt(16)
    heading4.runs[0].font.name = 'Calibri'
    
    if user_breakdown_df is not None and not user_breakdown_df.empty:
        # Add key metrics summary
        metrics_heading = doc.add_heading('4.1 Key Metrics Summary', level=2)
        metrics_heading.runs[0].font.size = Pt(14)
        metrics_heading.runs[0].font.name = 'Calibri'
        
        total_new_users = user_breakdown_df['new_users'].sum()
        total_returning_users = user_breakdown_df['old_users'].sum()
        total_all_users = user_breakdown_df['all_users'].sum()
        new_user_percentage = (total_new_users / total_all_users) * 100 if total_all_users > 0 else 0
        avg_new_users_per_month = user_breakdown_df['new_users'].mean()
        avg_returning_users_per_month = user_breakdown_df['old_users'].mean()
        
        # Create metrics table
        metrics_data = [
            ['Metric', 'Value'],
            ['Total New Users', f"{total_new_users:,}"],
            ['Total Returning Users', f"{total_returning_users:,}"],
            ['Total Users', f"{total_all_users:,}"],
            ['New User Percentage', f"{new_user_percentage:.1f}%"],
            ['Avg New Users/Month', f"{avg_new_users_per_month:.0f}"],
            ['Avg Returning Users/Month', f"{avg_returning_users_per_month:.0f}"]
        ]
        create_word_table(doc, metrics_data, '')
        
        # Add user breakdown chart
        chart_heading = doc.add_heading('4.2 User Breakdown Visualization', level=2)
        chart_heading.runs[0].font.size = Pt(14)
        chart_heading.runs[0].font.name = 'Calibri'
        
        try:
            # Generate the user breakdown chart
            chart_fig = create_user_breakdown_chart(user_breakdown_df, "All")
            
            # Create temporary image file for the chart
            chart_temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            chart_temp_file.close()
            
            # Save chart as image
            chart_fig.write_image(chart_temp_file.name, width=800, height=500, scale=2)
            
            # Add chart to document
            chart_para = doc.add_paragraph()
            chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = chart_para.runs[0] if chart_para.runs else chart_para.add_run()
            run.add_picture(chart_temp_file.name, width=Inches(6.5))
            
            # Add space after chart
            doc.add_paragraph()
            
            # Clean up temporary chart file
            os.unlink(chart_temp_file.name)
            
        except Exception as e:
            # Provide detailed error message
            if 'kaleido' in str(e).lower():
                error_msg = f"Chart generation unavailable: Kaleido package error - {str(e)}. Please ensure kaleido is installed: pip install --upgrade kaleido"
            else:
                error_msg = f"Chart generation unavailable: {str(e)}"
            
            chart_error = doc.add_paragraph(error_msg)
            chart_error.style = 'Normal'
        
        # Add user breakdown data table
        table_heading = doc.add_heading('4.3 Detailed User Breakdown Data', level=2)
        table_heading.runs[0].font.size = Pt(14)
        table_heading.runs[0].font.name = 'Calibri'
        
        # Prepare table data
        table_data = [['Month', 'Total Users', 'New Users', 'Returning Users', 'New User %', 'Returning User %']]
        
        for _, row in user_breakdown_df.iterrows():
            month_str = row['month'].strftime('%Y-%m-%d') if hasattr(row['month'], 'strftime') else str(row['month'])
            total_users = int(row['all_users'])
            new_users = int(row['new_users'])
            returning_users = int(row['old_users'])
            new_pct = (new_users / total_users * 100) if total_users > 0 else 0
            returning_pct = (returning_users / total_users * 100) if total_users > 0 else 0
            
            table_data.append([
                month_str,
                f"{total_users:,}",
                f"{new_users:,}",
                f"{returning_users:,}",
                f"{new_pct:.1f}%",
                f"{returning_pct:.1f}%"
            ])
        
        create_word_table(doc, table_data, '')
        
        # Analysis insights
        if len(user_breakdown_df) >= 2:
            insights_heading = doc.add_heading('4.4 Analysis Insights', level=2)
            insights_heading.runs[0].font.size = Pt(14)
            insights_heading.runs[0].font.name = 'Calibri'
            
            recent_new_users = user_breakdown_df.tail(3)['new_users'].mean()
            if len(user_breakdown_df) > 3:
                earlier_new_users = user_breakdown_df.head(len(user_breakdown_df)-3)['new_users'].mean()
            else:
                # Safety check: make sure we have at least 1 row
                if len(user_breakdown_df) >= 1:
                    earlier_new_users = user_breakdown_df.head(1)['new_users'].iloc[0]
                else:
                    earlier_new_users = 0
            
            total_months = len(user_breakdown_df)
            months_with_growth = len(user_breakdown_df[user_breakdown_df['new_users'] > user_breakdown_df['old_users']])
            if len(user_breakdown_df) > 0:
                best_month = user_breakdown_df.loc[user_breakdown_df['new_users'].idxmax()]
            else:
                best_month = {'month': 'N/A', 'new_users': 0}
            
            insights_text = f"""
• Analysis covers {total_months} months of user acquisition data
• {months_with_growth} out of {total_months} months showed more new users than returning users
• Customer base composition: {new_user_percentage:.1f}% new users, {(100-new_user_percentage):.1f}% returning users
• Best acquisition month: {best_month['month'].strftime('%Y-%m') if hasattr(best_month['month'], 'strftime') else str(best_month['month'])} with {best_month['new_users']:.0f} new users
"""
            
            if earlier_new_users > 0:
                new_user_trend = ((recent_new_users - earlier_new_users) / earlier_new_users) * 100
                trend_direction = "increasing" if new_user_trend > 5 else "decreasing" if new_user_trend < -5 else "stable"
                insights_text += f"• New user acquisition trend: {trend_direction} ({new_user_trend:+.1f}%)"
            
            insights_para = doc.add_paragraph(insights_text)
            insights_para.style = 'Normal'
            for run in insights_para.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
        
        # Add AI analysis if available
        if client:
            try:
                ai_heading = doc.add_heading('4.5 AI-Generated Analysis', level=2)
                ai_heading.runs[0].font.size = Pt(14)
                ai_heading.runs[0].font.name = 'Calibri'
                
                user_analysis = generate_user_breakdown_explanation(user_breakdown_df, model)
                if user_analysis['analysis']:
                    cleaned_analysis = clean_text_for_word(user_analysis['analysis'])
                    user_para = doc.add_paragraph(cleaned_analysis)
                    user_para.style = 'Normal'
                    for run in user_para.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Calibri'
            except Exception as e:
                error_para = doc.add_paragraph(f"AI user breakdown analysis unavailable: {str(e)}")
                error_para.style = 'Normal'
        else:
            ai_placeholder = doc.add_paragraph("AI-powered user breakdown analysis is available when an OpenAI API key is provided.")
            ai_placeholder.style = 'Normal'
    else:
        placeholder_para = doc.add_paragraph("User breakdown analysis is unavailable - no user breakdown data provided.")
        placeholder_para.style = 'Normal'
    
    # Save document
    doc.save(temp_file.name)
    
    return temp_file.name

def generate_pdf_report(cohort_table, selected_sku, raw_data, user_breakdown_df):
    """Generate a comprehensive PDF report with LTV analysis"""
    
    # Create a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    doc = SimpleDocTemplate(temp_file.name, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#1f77b4')
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12,
        spaceBefore=20,
        textColor=colors.HexColor('#2c3e50')
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=15,
        textColor=colors.HexColor('#34495e')
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor('#2c3e50')
    )
    
    # Title
    elements.append(Paragraph("Customer Lifetime Value (LTV) Analysis Report", title_style))
    elements.append(Spacer(1, 20))
    
    # Date and SKU info
    elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", body_style))
    elements.append(Paragraph(f"Analysis Scope: {selected_sku if selected_sku != 'All' else 'All Products'}", body_style))
    elements.append(Spacer(1, 30))
    
    # Executive Summary
    elements.append(Paragraph("Executive Summary", heading_style))
    
    # Calculate key metrics from cohort data
    ltv_data = cohort_table[cohort_table['Metric'] == 'LTV'].copy()
    revenue_data = cohort_table[cohort_table['Metric'] == 'Revenue'].copy()
    
    if not ltv_data.empty:
        month_columns = [col for col in ltv_data.columns if col not in ['POME Month', 'Cohort Size', 'Metric', 'Total']]
        month_columns = sorted(month_columns)
        
        # Calculate LTV performance metrics
        first_month_ltv_values = []
        latest_ltv_values = []
        total_customers = 0
        
        for _, row in ltv_data.iterrows():
            pome_month = row['POME Month']
            cohort_size = row['Cohort Size']
            total_customers += cohort_size
            
            if pome_month in month_columns:
                first_ltv_str = row[pome_month]
                if isinstance(first_ltv_str, str) and first_ltv_str.startswith('$'):
                    first_ltv = float(first_ltv_str.replace('$', '').replace(',', ''))
                    if first_ltv > 0:
                        first_month_ltv_values.append(first_ltv)
                        
                        # Get latest month LTV
                        latest_ltv = 0
                        for month_col in reversed(month_columns):
                            if month_col >= pome_month:
                                ltv_str = row[month_col]
                                if isinstance(ltv_str, str) and ltv_str.startswith('$'):
                                    ltv_val = float(ltv_str.replace('$', '').replace(',', ''))
                                    if ltv_val > 0:
                                        latest_ltv = ltv_val
                                        break
                        latest_ltv_values.append(latest_ltv)
        
        if first_month_ltv_values and latest_ltv_values:
            avg_first_ltv = sum(first_month_ltv_values) / len(first_month_ltv_values)
            avg_latest_ltv = sum(latest_ltv_values) / len(latest_ltv_values)
            ltv_multiplier = avg_latest_ltv / avg_first_ltv if avg_first_ltv > 0 else 0
                        
            elements.append(Spacer(1, 20))
    
    # Build PDF
    doc.build(elements)
    temp_file.close()
    
    return temp_file.name

def create_download_link(df, filename, file_label):
    """Create a download link for a dataframe"""
    csv = df.to_csv(index=False)
    b64 = io.StringIO(csv).getvalue().encode()
    href = f'<a href="data:file/csv;base64,{b64.hex()}" download="{filename}" style="text-decoration: none;"><button style="background-color: #4CAF50; color: white; padding: 8px 16px; border: none; border-radius: 4px; cursor: pointer;">{file_label}</button></a>'
    return href

# Main Streamlit App
def main():
    st.title("📊 LTV Calculation Dashboard")
    st.markdown("Upload your CSV files and generate ProductRaw, ProductSummary, RawData, and User Lifecycle Analysis")
    
    # Sidebar for file upload
    st.sidebar.header("📁 File Upload")
    uploaded_files = st.sidebar.file_uploader(
        "Choose CSV files",
        type="csv",
        accept_multiple_files=True,
        help="Upload multiple CSV files to combine and analyze"
    )
    
    if uploaded_files:
        st.sidebar.success(f"✅ {len(uploaded_files)} files uploaded")
        
        # Process files
        with st.spinner("Processing uploaded files..."):
            combined_df = process_files(uploaded_files)
        
        if combined_df is not None:
            st.success("Files processed successfully!")
            
            # Display basic info
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Total Rows", f"{len(combined_df):,}")
            with col2:
                st.metric("Total Columns", len(combined_df.columns))
            with col3:
                st.metric("Files Processed", len(uploaded_files))
            
            # Currency Selection
            st.header("💰 Currency Selection")
            available_currencies = sorted(combined_df['currency'].unique())
            selected_currency = st.selectbox(
                "Select Currency",
                options=available_currencies,
                index=0 if 'USD' not in available_currencies else available_currencies.index('USD'),
                help="Choose the currency for analysis. All calculations will be filtered by this currency."
            )
            st.info(f"Selected Currency: **{selected_currency}**")
            
            # Calculate all dataframes
            st.header("📈 Calculations")
            
            with st.spinner("Calculating ProductRaw..."):
                product_raw = calculate_product_raw(combined_df, selected_currency)
            
            with st.spinner("Calculating ProductSummary..."):
                product_summary = calculate_product_summary(combined_df, selected_currency)
            
            with st.spinner("Calculating RawData..."):
                raw_data = calculate_raw_data(combined_df, selected_currency)

            with st.spinner("Calculating RawDataWithoutSku..."):
                raw_data_wo_sku = calculate_raw_data_wo_sku(combined_df, selected_currency)
            
            st.success("All calculations completed!")
            
            # Display results in tabs
            tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs(["ProductRaw", "ProductSummary", "RawData", "User Lifecycle Analysis", "Retention & LTV Analysis", "User Breakdown Analysis", "📊 Product LTV Analysis", "📄 Top 10 Products Report"])
            
            with tab1:
                st.subheader("ProductRaw Data")
                st.dataframe(product_raw, use_container_width=True)
                st.markdown(create_download_link(product_raw, "ProductRaw.csv", "📥 Download ProductRaw"), unsafe_allow_html=True)
            
            with tab2:
                st.subheader("ProductSummary Data")
                st.dataframe(product_summary, use_container_width=True)
                st.markdown(create_download_link(product_summary, "ProductSummary.csv", "📥 Download ProductSummary"), unsafe_allow_html=True)
            
            with tab3:
                st.subheader("RawData")
                st.dataframe(raw_data, use_container_width=True)
                st.markdown(create_download_link(raw_data, "RawData.csv", "📥 Download RawData"), unsafe_allow_html=True)
            
            with tab4:
                st.subheader("User Lifecycle Analysis")
                
                # Merchant SKU selection
                available_skus = ["All"] + sorted(raw_data['merchant_sku'].unique().tolist())
                selected_sku = st.selectbox(
                    "Select Merchant SKU for User Lifecycle Analysis:",
                    available_skus,
                    help="Choose a specific SKU or 'All' to analyze all SKUs together"
                )
                
                # Calculate cohort analysis
                with st.spinner(f"Calculating User Lifecycle Analysis for {selected_sku}..."):
                    cohort_table, filter_msg = calculate_cohort_analysis(raw_data, selected_sku)
                
                st.info(filter_msg)
                
                # Display cohort table with scrolling
                st.subheader("User Lifecycle Analysis Results")
                st.dataframe(
                    cohort_table, 
                    use_container_width=True,
                    height=600  # Make it scrollable
                )
                
                # Download button for cohort analysis
                st.markdown(create_download_link(cohort_table, f"Cohort_Analysis_{selected_sku}.csv", "📥 Download User Lifecycle Analysis"), unsafe_allow_html=True)
            
            with tab5:
                st.header("📈 Retention & LTV Analysis")
                st.markdown("""
                This tab contains complementary analyses:
                 **Cumulative LTV Analysis** - Cumulative lifetime value from cohort data
                """)
                
                # # MoM Retention Analysis Section
                # st.subheader("📊 MoM Retention Analysis (ARPU)")
                # st.markdown("""
                # **Month-over-Month Retention Analysis** shows the Average Revenue Per User (ARPU) for each cohort across different months.
                
                # This analysis uses the Excel formula: 
                # `=IF(D$3<$B14,"",SUMIFS(sales,pome_month,cohort_month,month,analysis_month)/SUMIFS(users,pome_month,cohort_month,month,analysis_month))`
                # """)
                
                # # Calculate retention analysis
                # with st.spinner("Calculating MoM Retention..."):
                #     try:
                #         retention_table = retention_calculation(raw_data, cohort_table)
                        
                #         if retention_table is not None and not retention_table.empty:
                #             st.success("✅ MoM Retention calculated successfully!")
                            
                #             # # Display key metrics
                #             # col1, col2, col3 = st.columns(3)
                #             # with col1:
                #             #     total_cohorts = len(retention_table)
                #             #     st.metric("Total Cohorts", total_cohorts)
                #             # with col2:
                #             #     avg_cohort_size = retention_table['Cohort Size'].mean()
                #             #     st.metric("Avg Cohort Size", f"{avg_cohort_size:.0f}")
                #             # with col3:
                #             #     total_customers = retention_table['Cohort Size'].sum()
                #             #     st.metric("Total Customers", f"{total_customers:,}")
                            
                #             # Display retention table
                #             st.subheader("MoM Retention Table (ARPU by Cohort)")
                #             st.info("💡 Values represent Average Revenue Per User (ARPU) for each cohort in each month")
                            
                #             # Format the display table for better readability
                #             display_table = retention_table.copy()
                            
                #             # Format numeric columns (skip POME Month and Cohort Size)
                #             numeric_cols = [col for col in display_table.columns if col not in ['POME Month']]
                #             for col in numeric_cols:
                #                 display_table[col] = display_table[col].apply(
                #                     lambda x: f"${x:.2f}" if isinstance(x, (int, float)) and x != 0 else (x if x != 0 else "")
                #                 )
                            
                #             # Display with styling
                #             st.dataframe(
                #                 display_table,
                #                 use_container_width=True,
                #                 height=600
                #             )
                            
                #             # Download button
                #             st.markdown(create_download_link(retention_table, "MoM_Retention_Analysis.csv", "📥 Download MoM Retention"), unsafe_allow_html=True)
                            
                #             # Analysis insights
                #             st.subheader("📊 Key Insights")
                            
                #             # Calculate some insights
                #             numeric_retention_data = retention_table.copy()
                #             for col in numeric_cols:
                #                 numeric_retention_data[col] = pd.to_numeric(numeric_retention_data[col], errors='coerce')
                            
                #             # Find highest ARPU
                #             arpu_values = []
                #             for col in numeric_cols:
                #                 col_values = numeric_retention_data[col].dropna()
                #                 if len(col_values) > 0:
                #                     arpu_values.extend(col_values.tolist())
                            
                #             if arpu_values:
                #                 max_arpu = max([x for x in arpu_values if x > 0])
                #                 avg_arpu = np.mean([x for x in arpu_values if x > 0])
                                
                #                 col1, col2 = st.columns(2)
                #                 with col1:
                #                     st.metric("Highest ARPU", f"${max_arpu:.2f}")
                #                 with col2:
                #                     st.metric("Average ARPU", f"${avg_arpu:.2f}")
                #         else:
                #             st.error("❌ Failed to calculate retention analysis. Please check your data.")
                            
                #     except Exception as e:
                #         st.error(f"❌ Error calculating retention: {str(e)}")
                #         st.error("Please ensure your cohort analysis was calculated successfully first.")
                
                # # Add separator
                # st.markdown("---")
                
                # Cumulative LTV Analysis Section
                st.subheader("💰 Cumulative LTV Analysis")
                st.markdown("""
                **Cumulative Lifetime Value Analysis** shows the cumulative LTV for each cohort across different months. 
                This is calculated by taking the LTV values from the User Lifecycle Analysis and showing them cumulatively from lowest month to highest month.
                
                💡 **How it works**: For each cohort, the LTV values are summed cumulatively across months, showing the total lifetime value accumulated over time.
                """)
                
                # Calculate cumulative LTV analysis
                with st.spinner("Calculating Cumulative LTV Analysis..."):
                    try:
                        cumulative_ltv_table = retention_calculation_v2(raw_data, cohort_table)
                        
                        if cumulative_ltv_table is not None and not cumulative_ltv_table.empty:
                            st.success("✅ Cumulative LTV Analysis calculated successfully!")
                            
                            # Display metrics for cumulative LTV
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                total_cohorts_ltv = len(cumulative_ltv_table)
                                st.metric("Total Cohorts", total_cohorts_ltv)
                            with col2:
                                # Find the highest cumulative LTV value
                                numeric_cols_ltv = [col for col in cumulative_ltv_table.columns if col not in ['POME Month', 'Cohort Size']]
                                max_ltv_values = []
                                for _, row in cumulative_ltv_table.iterrows():
                                    row_values = [row[col] for col in numeric_cols_ltv if isinstance(row[col], (int, float)) and row[col] > 0]
                                    if row_values:
                                        max_ltv_values.append(max(row_values))
                                
                                if max_ltv_values:
                                    highest_ltv = max(max_ltv_values)
                                    st.metric("Highest Cumulative LTV", f"${highest_ltv:.2f}")
                                else:
                                    st.metric("Highest Cumulative LTV", "$0.00")
                            with col3:
                                # Calculate average final LTV (last non-empty value for each cohort)
                                final_ltv_values = []
                                for _, row in cumulative_ltv_table.iterrows():
                                    row_values = [row[col] for col in reversed(numeric_cols_ltv) if isinstance(row[col], (int, float)) and row[col] > 0]
                                    if row_values:
                                        final_ltv_values.append(row_values[0])  # First non-zero from the right (latest month)
                                
                                if final_ltv_values:
                                    avg_final_ltv = sum(final_ltv_values) / len(final_ltv_values)
                                    st.metric("Avg Final LTV", f"${avg_final_ltv:.2f}")
                                else:
                                    st.metric("Avg Final LTV", "$0.00")
                            
                            # Display cumulative LTV table
                            st.subheader("Cumulative LTV Table by Cohort")
                            st.info("💡 Values represent cumulative LTV for each cohort across months. Values increase from left to right showing total lifetime value accumulation.")
                            
                            # Format the display table for better readability
                            display_table_ltv = cumulative_ltv_table.copy()
                            
                            # Format numeric columns (skip POME Month and Cohort Size)
                            numeric_cols_ltv = [col for col in display_table_ltv.columns if col not in ['POME Month', 'Cohort Size']]
                            for col in numeric_cols_ltv:
                                display_table_ltv[col] = display_table_ltv[col].apply(
                                    lambda x: f"${x:.2f}" if isinstance(x, (int, float)) and x != 0 else (x if x != 0 else "")
                                )
                            
                            # Display with styling
                            st.dataframe(
                                display_table_ltv,
                                use_container_width=True,
                                height=600
                            )
                            
                            # Download button for cumulative LTV
                            st.markdown(create_download_link(cumulative_ltv_table, "Cumulative_LTV_Analysis.csv", "📥 Download Cumulative LTV"), unsafe_allow_html=True)
                            
                        else:
                            st.error("❌ Failed to calculate cumulative LTV analysis. Please ensure your User Lifecycle Analysis contains LTV data.")
                            
                    except Exception as e:
                        st.error(f"❌ Error calculating cumulative LTV: {str(e)}")
                        st.error("Please ensure your User Lifecycle Analysis was calculated successfully and contains LTV metrics.")
            
            with tab6:
                st.header("👥 User Breakdown Analysis")
                st.markdown("""
                **New vs Returning Users Analysis** provides insights into customer acquisition and retention patterns by showing the breakdown of users each month.
                
                This analysis helps answer key business questions:
                - How many new customers are we acquiring each month?
                - What's the ratio of new vs returning customers?
                - How does user composition change over time?
                """)
                
                # Merchant SKU selection for user breakdown
                st.subheader("🎯 Analysis Configuration")
                available_skus = ["All"] + sorted(raw_data['merchant_sku'].unique().tolist())
                selected_sku_breakdown = st.selectbox(
                    "Select Merchant SKU for User Breakdown:",
                    available_skus,
                    help="Choose a specific SKU or 'All' to analyze user patterns across all SKUs",
                    key="user_breakdown_sku_selector"
                )
                
                st.markdown("---")
                
                # User breakdown explanation
                st.subheader("📖 Methodology")
                st.markdown("""
                **Calculation Logic**:
                - **New Users**: Users where their first purchase month (POME) equals the analysis month
                - **Returning Users**: Total users minus new users for each month
                - **Total Users**: All unique users who made purchases in each month
                
                This segmentation helps identify customer lifecycle patterns and acquisition effectiveness.
                """)
                
                # Calculate user breakdown
                with st.spinner(f"Calculating user breakdown for {selected_sku_breakdown}..."):
                    user_breakdown = calculate_user_breakdown(raw_data, raw_data_wo_sku, selected_sku_breakdown)
                
                if not user_breakdown.empty:
                    # Display key metrics
                    st.subheader("📊 Key Metrics")
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        total_new_users = user_breakdown['new_users'].sum()
                        st.metric("Total New Users", f"{total_new_users:,}")
                    with col2:
                        total_returning_users = user_breakdown['old_users'].sum()
                        st.metric("Total Returning Users", f"{total_returning_users:,}")
                    with col3:
                        total_all_users = user_breakdown['all_users'].sum()
                        st.metric("Total Users", f"{total_all_users:,}")
                    with col4:
                        if total_all_users > 0:
                            new_user_percentage = (total_new_users / total_all_users) * 100
                            st.metric("New User %", f"{new_user_percentage:.1f}%")
                        else:
                            st.metric("New User %", "0%")
                    
                    # Additional insights
                    col1, col2 = st.columns(2)
                    with col1:
                        avg_new_users_per_month = user_breakdown['new_users'].mean()
                        st.metric("Avg New Users/Month", f"{avg_new_users_per_month:.0f}")
                    with col2:
                        avg_returning_users_per_month = user_breakdown['old_users'].mean()
                        st.metric("Avg Returning Users/Month", f"{avg_returning_users_per_month:.0f}")
                    
                    # Create and display the stacked bar chart
                    st.subheader("📈 User Breakdown Visualization")
                    fig = create_user_breakdown_chart(user_breakdown, selected_sku_breakdown)
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Additional analysis insights
                    st.subheader("🔍 Analysis Insights")
                    
                    # Calculate trends
                    if len(user_breakdown) >= 2:
                        recent_new_users = user_breakdown.tail(3)['new_users'].mean()
                        if len(user_breakdown) > 3:
                            earlier_new_users = user_breakdown.head(len(user_breakdown)-3)['new_users'].mean()
                        else:
                            # Safety check: make sure we have at least 1 row
                            if len(user_breakdown) >= 1:
                                earlier_new_users = user_breakdown.head(1)['new_users'].iloc[0]
                            else:
                                earlier_new_users = 0
                        
                        if earlier_new_users > 0:
                            new_user_trend = ((recent_new_users - earlier_new_users) / earlier_new_users) * 100
                            trend_direction = "📈 Increasing" if new_user_trend > 5 else "📉 Decreasing" if new_user_trend < -5 else "➡️ Stable"
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                st.info(f"**New User Acquisition Trend**: {trend_direction} ({new_user_trend:+.1f}%)")
                            with col2:
                                best_month = user_breakdown.loc[user_breakdown['new_users'].idxmax()]
                                st.info(f"**Best Acquisition Month**: {best_month['month'].strftime('%Y-%m')} ({best_month['new_users']:.0f} new users)")
                    
                    # Display data table
                    st.subheader("📋 Detailed User Breakdown Data")
                    display_breakdown = user_breakdown.copy()
                    display_breakdown['month'] = display_breakdown['month'].dt.strftime('%Y-%m-%d')
                    display_breakdown = display_breakdown.rename(columns={
                        'month': 'Month',
                        'all_users': 'Total Users',
                        'new_users': 'New Users',
                        'old_users': 'Returning Users'
                    })
                    
                    # Add percentage columns
                    display_breakdown['New User %'] = ((display_breakdown['New Users'] / display_breakdown['Total Users']) * 100).round(1)
                    display_breakdown['Returning User %'] = ((display_breakdown['Returning Users'] / display_breakdown['Total Users']) * 100).round(1)
                    
                    st.dataframe(
                        display_breakdown,
                        use_container_width=True,
                        height=400
                    )
                    
                    # Download button for user breakdown
                    st.markdown(create_download_link(display_breakdown, f"User_Breakdown_{selected_sku_breakdown}.csv", "📥 Download User Breakdown Analysis"), unsafe_allow_html=True)
                    
                    # Summary insights
                    st.subheader("💡 Summary Insights")
                    total_months = len(user_breakdown)
                    months_with_growth = len(user_breakdown[user_breakdown['new_users'] > user_breakdown['old_users']])
                    
                    insights = [
                        f"📅 **Analysis Period**: {total_months} months of data",
                        f"🆕 **Acquisition-Heavy Months**: {months_with_growth} out of {total_months} months had more new users than returning users",
                        f"📊 **Customer Base Composition**: {new_user_percentage:.1f}% new users, {(100-new_user_percentage):.1f}% returning users"
                    ]
                    
                    for insight in insights:
                        st.markdown(insight)
                
                else:
                    st.warning("No user breakdown data available for the selected criteria.")
                    st.info("Please ensure you have processed data with valid purchase dates and user information.")
            
            with tab7:
                st.header("📊 Product LTV Analysis")
                st.markdown("""
                **Product-level LTV Analysis** provides detailed insights into the lifetime value performance of each product (merchant SKU).
                This analysis helps answer key questions:
                - Which products drive the highest customer lifetime value?
                - How do different products perform in terms of retention and repeat purchases?
                - What are the monthly progression patterns for each product?
                """)
                
                # Product LTV Analysis Section
                st.subheader("🎯 Product Performance Analysis")
                
                with st.spinner("Calculating Product LTV Analysis..."):
                    try:
                        # Create product LTV table
                        product_ltv_data = create_product_ltv_table(product_raw, raw_data)
                        
                        if product_ltv_data:
                            st.success(f"✅ Product LTV Analysis completed for {len(product_ltv_data)} top products!")
                            
                            # Display key metrics
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                total_products_analyzed = len(product_ltv_data)
                                st.metric("Products Analyzed", total_products_analyzed)
                            with col2:
                                total_cohort_size = sum([product['Cohort Size'] for product in product_ltv_data])
                                st.metric("Total Customers", f"{total_cohort_size:,}")
                            with col3:
                                avg_cohort_size = total_cohort_size / total_products_analyzed if total_products_analyzed > 0 else 0
                                st.metric("Avg Cohort Size", f"{avg_cohort_size:.0f}")
                            
                            # Display the complete Product LTV data
                            st.subheader("📊 Complete Product LTV Analysis")
                            st.markdown("**Detailed monthly metrics for all analyzed products**")
                            
                            # Generate the exportable table for display
                            display_df = export_product_ltv_table(product_ltv_data)
                            
                            # Show the complete table
                            st.dataframe(display_df, use_container_width=True, height=600)
                            
                            # Export functionality
                            st.subheader("📥 Export Product LTV Analysis")
                            
                            # Generate the exportable table
                            export_df = export_product_ltv_table(product_ltv_data)
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                st.info(f"📊 Export contains {len(export_df)} rows with detailed monthly metrics for each product")
                            with col2:
                                st.info("📋 Includes: Active Customers, Purchases, Revenue, Cumulative Revenue, Retention Rate, LTV")
                            
                            # Download button for the detailed export
                            st.markdown(create_download_link(export_df, "Product_LTV_Analysis.csv", "📥 Download Product LTV Analysis"), unsafe_allow_html=True)
                            
                            # Show preview of export structure
                            with st.expander("👁️ Preview Export Structure"):
                                st.markdown("**First 20 rows of the export table:**")
                                st.dataframe(export_df.head(20), use_container_width=True)
                            
                            # Add the Top Products Dashboard
                            st.markdown("---")
                            st.subheader("📊 Top Products Dashboard")
                            st.markdown("**Key performance metrics across all products in your portfolio**")
                            
                            # Calculate the top products tables
                            top_products_data = calculate_top_products_tables(product_raw, raw_data)
                            
                            # Display in 2x2 grid
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.markdown("##### 🎯 Top 10 Products by Acquired Customers")
                                acquired_df = pd.DataFrame(top_products_data['top_acquired'])
                                # Clean numeric data to avoid PyArrow conversion errors
                                for col in acquired_df.columns:
                                    if col not in ['Product Title', 'Merchant SKU']:
                                        if acquired_df[col].dtype == 'object':
                                            acquired_df[col] = pd.to_numeric(acquired_df[col], errors='coerce').fillna(0)
                                # Add ranking
                                acquired_df.insert(0, 'Rank', range(1, len(acquired_df) + 1))
                                st.dataframe(acquired_df, use_container_width=True, height=350)
                                
                            with col2:
                                st.markdown("##### 🔄 Top 10 Products by Repeat Rate")
                                repeat_df = pd.DataFrame(top_products_data['top_repeat'])
                                # Clean numeric data to avoid PyArrow conversion errors
                                for col in repeat_df.columns:
                                    if col not in ['Product Title', 'Merchant SKU']:
                                        if repeat_df[col].dtype == 'object':
                                            # Special handling for percentage strings
                                            if 'Rate' in col or '%' in str(repeat_df[col].iloc[0]):
                                                repeat_df[col] = repeat_df[col].astype(str).str.replace('%', '').replace('', '0')
                                                repeat_df[col] = pd.to_numeric(repeat_df[col], errors='coerce').fillna(0)
                                            else:
                                                repeat_df[col] = pd.to_numeric(repeat_df[col], errors='coerce').fillna(0)
                                # Add ranking
                                repeat_df.insert(0, 'Rank', range(1, len(repeat_df) + 1))
                                st.dataframe(repeat_df, use_container_width=True, height=350)
                            
                            col3, col4 = st.columns(2)
                            
                            with col3:
                                st.markdown("##### 💰 Top 10 Products by AOV")
                                aov_df = pd.DataFrame(top_products_data['top_aov'])
                                # Clean numeric data to avoid PyArrow conversion errors
                                for col in aov_df.columns:
                                    if col not in ['Product Title', 'Merchant SKU']:
                                        if aov_df[col].dtype == 'object':
                                            # Special handling for currency strings
                                            if 'AOV' in col or '$' in str(aov_df[col].iloc[0]):
                                                aov_df[col] = aov_df[col].astype(str).str.replace('$', '').replace('', '0')
                                                aov_df[col] = pd.to_numeric(aov_df[col], errors='coerce').fillna(0)
                                            else:
                                                aov_df[col] = pd.to_numeric(aov_df[col], errors='coerce').fillna(0)
                                # Add ranking
                                aov_df.insert(0, 'Rank', range(1, len(aov_df) + 1))
                                st.dataframe(aov_df, use_container_width=True, height=350)
                                
                            with col4:
                                st.markdown("##### 📈 Top 10 Products by LTV")
                                ltv_df = pd.DataFrame(top_products_data['top_ltv'])
                                # Clean numeric data to avoid PyArrow conversion errors
                                for col in ltv_df.columns:
                                    if col not in ['Product Title', 'Merchant SKU']:
                                        if ltv_df[col].dtype == 'object':
                                            # Special handling for currency strings
                                            if 'LTV' in col or '$' in str(ltv_df[col].iloc[0]):
                                                ltv_df[col] = ltv_df[col].astype(str).str.replace('$', '').replace('', '0')
                                                ltv_df[col] = pd.to_numeric(ltv_df[col], errors='coerce').fillna(0)
                                            else:
                                                ltv_df[col] = pd.to_numeric(ltv_df[col], errors='coerce').fillna(0)
                                # Add ranking
                                ltv_df.insert(0, 'Rank', range(1, len(ltv_df) + 1))
                                st.dataframe(ltv_df, use_container_width=True, height=350)
                            
                            # Download all top products tables
                            st.subheader("📥 Download Top Products Analysis")
                            
                            # Create a combined Excel-style export
                            col1, col2, col3, col4 = st.columns(4)
                            
                            with col1:
                                st.markdown(create_download_link(acquired_df, "Top_Products_Acquired_Customers.csv", "📥 Acquired Customers"), unsafe_allow_html=True)
                            with col2:
                                st.markdown(create_download_link(repeat_df, "Top_Products_Repeat_Rate.csv", "📥 Repeat Rate"), unsafe_allow_html=True)
                            with col3:
                                st.markdown(create_download_link(aov_df, "Top_Products_AOV.csv", "📥 AOV"), unsafe_allow_html=True)
                            with col4:
                                st.markdown(create_download_link(ltv_df, "Top_Products_LTV.csv", "📥 LTV"), unsafe_allow_html=True)
                            

                        else:
                            st.error("❌ Failed to generate Product LTV Analysis. Please ensure ProductRaw data is available.")
                            
                    except Exception as e:
                        st.error(f"❌ Error calculating Product LTV Analysis: {str(e)}")
                        st.error("Please ensure ProductRaw and RawData have been calculated successfully.")
            
            with tab8:
                st.header("📄 Comprehensive Report")
                st.markdown("**Generate comprehensive Word report with cohort analysis, user breakdown tables, and AI-powered LTV insights for all top 10 products**")
                
                # OpenAI Configuration Section
                st.subheader("🤖 AI-Powered Analysis Configuration")
                
                col1, col2 = st.columns([2, 1])
                with col1:
                    openai_api_key = st.text_input(
                        "OpenAI API Key",
                        type="password",
                        help="Enter your OpenAI API key to enable AI-powered LTV analysis for each product",
                        placeholder="sk-..."
                    )
                with col2:
                    model_choice = st.selectbox(
                        "Model",
                        ["gpt-4o-mini", "gpt-4o", "gpt-3.5-turbo"],
                        help="Choose the OpenAI model for analysis"
                    )
                
                # Initialize OpenAI client if API key is provided
                global client
                if openai_api_key:
                    try:
                        # Clean the API key to remove any chat formatting or Unicode characters
                        cleaned_api_key = clean_api_key(openai_api_key)
                        client = OpenAI(api_key=cleaned_api_key)
                        st.success("✅ AI analysis enabled! Your report will include expert LTV insights for each product.")
                    except Exception as e:
                        st.error(f"❌ Error initializing OpenAI client: {str(e)}")
                        client = None
                else:
                    client = None
                    st.info("💡 **Optional**: Provide an OpenAI API key to get AI-powered LTV analysis and business insights for each product in your report.")
                
                st.markdown("---")
                
                st.markdown("**Comprehensive LTV Report Includes:**")
                st.markdown("""
                • **🏢 Account-Level LTV Analysis** - Overall business performance across all products with cohort walk-throughs and benchmark comparisons
                • **🏆 Top 10 Product Performance Tables** - Rankings by Acquired Customers, Repeat Rate, AOV, and LTV with strategic insights
                • **🔍 Individual Product Analysis** - Detailed LTV analysis for each of the top 10 products by acquired customers
                • **👥 User Acquisition & Retention Analysis** - New vs returning customer patterns and business health indicators
                • **🤖 AI-Powered Insights** - Expert recommendations and strategic guidance throughout (when API key provided)
                • **📊 Professional Formatting** - Executive-ready report in Word format optimized for presentations and decision-making
                """)
                
                # Generate comprehensive report button
                if st.button("🚀 Generate Comprehensive LTV Report", type="primary"):
                    with st.spinner("Generating comprehensive LTV report..."):
                        try:
                            # Calculate account-level cohort analysis and user breakdown for the comprehensive report
                            account_cohort_table, _ = calculate_cohort_analysis(raw_data, "All")
                            account_user_breakdown = calculate_user_breakdown(raw_data, raw_data_wo_sku, "All")
                            
                            # Generate the comprehensive Word report
                            word_path = generate_comprehensive_word_report(
                                product_raw, raw_data, account_cohort_table, account_user_breakdown, model_choice
                            )
                            
                            # Read the PDF file
                            # with open(pdf_path, "rb") as pdf_file:
                            #     pdf_bytes = pdf_file.read()
                                
                            # Read the Word file
                            with open(word_path, "rb") as word_file:
                                word_bytes = word_file.read()
                            
                            # Provide download buttons
                            ai_status = "with AI-powered insights" if client else "with data tables only"
                            st.success(f"✅ Comprehensive LTV report generated successfully {ai_status}!")
                            
                            # col1, col2 = st.columns(2)
                            # with col1:
                            #     st.download_button(
                            #         label="📥 Download PDF Report",
                            #         data=pdf_bytes,
                            #         file_name=f"Comprehensive_LTV_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            #         mime="application/pdf"
                            #     )
                            # with col2:
                            st.download_button(
                                label="📥 Download Word Report",
                                data=word_bytes,
                                file_name=f"Comprehensive_LTV_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            
                            # Clean up temporary files
                            # os.unlink(pdf_path)
                            os.unlink(word_path)
                            
                            # Show what was included
                            ai_info = "**🤖 AI Analysis:** Expert LTV insights and recommendations for each product" if client else "**📊 Analysis:** Placeholder sections for AI insights (API key required)"
                            st.info(f"""
                            **Report Generated:** Top 10 Products Analysis (Word format)  
                            **Products Included:** Top 10 products by acquired customers  
                            **Content:** AI-powered strategic insights with data file references  
                            {ai_info}  
                            **Model Used:** {model_choice if client else "N/A"}  
                            **Data Files:** Available for download separately below
                            """)
                            
                        except Exception as e:
                            st.error(f"❌ Error generating comprehensive report: {str(e)}")
                            st.info("Please ensure all data has been calculated successfully and try again.")
                            
                # Add download section for individual product data files
                st.markdown("---")
                st.subheader("📊 Download Individual Product Data Files")
                st.markdown("**Get detailed data files for each of the top 10 products (referenced in the Word report)**")
                
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.info("""
                    **Data Files Include:**
                    • **Cohort Analysis** - Monthly user behavior and LTV progression for each product
                    • **User Breakdown** - New vs returning customer analysis for each product  
                    • **Product LTV Data** - Detailed LTV metrics and calculations for each product
                    
                    These are the same data files referenced in your Word report for detailed analysis.
                    """)
                
                with col2:
                    if st.button("📥 Download Product Data Files", type="secondary"):
                        try:
                            with st.spinner("Generating data files for top 10 products..."):
                                # Generate the data files
                                zip_data = generate_product_data_files(product_raw, raw_data)
                                
                                # Provide download button
                                st.success("✅ Product data files generated successfully!")
                                st.download_button(
                                    label="📥 Download ZIP File",
                                    data=zip_data,
                                    file_name=f"Top_10_Products_Data_Files_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                                    mime="application/zip"
                                )
                                
                                st.info("""
                                **Generated Files:** Individual CSV files for each of the top 10 products  
                                **File Naming:** `[SKU]_[DataType].csv` (e.g., `ABC123_Cohort_Analysis.csv`)  
                                **Contents:** Detailed data tables referenced in your Word report
                                """)
                                
                        except Exception as e:
                            st.error(f"❌ Error generating data files: {str(e)}")
            
            # Download all results as ZIP
            st.header("📦 Download All Results")
            if st.button("📥 Download All as ZIP"):
                # Create ZIP file
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    # Add each DataFrame as CSV to ZIP
                    zip_file.writestr("ProductRaw.csv", product_raw.to_csv(index=False))
                    zip_file.writestr("ProductSummary.csv", product_summary.to_csv(index=False))
                    zip_file.writestr("RawData.csv", raw_data.to_csv(index=False))
                    zip_file.writestr(f"Cohort_Analysis_{selected_sku}.csv", cohort_table.to_csv(index=False))
                    
                    # Add retention analysis if available
                    try:
                        retention_table = retention_calculation(raw_data, cohort_table)
                        if retention_table is not None and not retention_table.empty:
                            zip_file.writestr("MoM_Retention_Analysis.csv", retention_table.to_csv(index=False))
                    except:
                        pass  # Skip if retention calculation fails
                    
                    # Add cumulative LTV analysis if available
                    try:
                        cumulative_ltv_table = retention_calculation_v2(raw_data, cohort_table)
                        if cumulative_ltv_table is not None and not cumulative_ltv_table.empty:
                            zip_file.writestr("Cumulative_LTV_Analysis.csv", cumulative_ltv_table.to_csv(index=False))
                    except:
                        pass  # Skip if cumulative LTV calculation fails
                    
                    # Add user breakdown analysis if available
                    try:
                        # Use the "All" option for ZIP download to include comprehensive data
                        user_breakdown = calculate_user_breakdown(raw_data, raw_data_wo_sku, "All")
                        if not user_breakdown.empty:
                            display_breakdown = user_breakdown.copy()
                            display_breakdown['month'] = display_breakdown['month'].dt.strftime('%Y-%m-%d')
                            display_breakdown = display_breakdown.rename(columns={
                                'month': 'Month',
                                'all_users': 'Total Users',
                                'new_users': 'New Users',
                                'old_users': 'Returning Users'
                            })
                            # Add percentage columns
                            display_breakdown['New User %'] = ((display_breakdown['New Users'] / display_breakdown['Total Users']) * 100).round(1)
                            display_breakdown['Returning User %'] = ((display_breakdown['Returning Users'] / display_breakdown['Total Users']) * 100).round(1)
                            zip_file.writestr("User_Breakdown_Analysis.csv", display_breakdown.to_csv(index=False))
                    except:
                        pass  # Skip if user breakdown calculation fails
                    
                    # Add Product LTV Analysis if available
                    try:
                        product_ltv_data = create_product_ltv_table(product_raw, raw_data)
                        if product_ltv_data:
                            product_ltv_export = export_product_ltv_table(product_ltv_data)
                            zip_file.writestr("Product_LTV_Analysis.csv", product_ltv_export.to_csv(index=False))
                    except:
                        pass  # Skip if Product LTV calculation fails
                    
                    # Add Top Products Tables if available
                    try:
                        top_products_data = calculate_top_products_tables(product_raw, raw_data)
                        
                        # Add each top products table with data cleaning
                        acquired_df = pd.DataFrame(top_products_data['top_acquired'])
                        for col in acquired_df.columns:
                            if col not in ['Product Title', 'Merchant SKU'] and acquired_df[col].dtype == 'object':
                                acquired_df[col] = pd.to_numeric(acquired_df[col], errors='coerce').fillna(0)
                        acquired_df.insert(0, 'Rank', range(1, len(acquired_df) + 1))
                        zip_file.writestr("Top_Products_Acquired_Customers.csv", acquired_df.to_csv(index=False))
                        
                        repeat_df = pd.DataFrame(top_products_data['top_repeat'])
                        for col in repeat_df.columns:
                            if col not in ['Product Title', 'Merchant SKU'] and repeat_df[col].dtype == 'object':
                                if 'Rate' in col or '%' in str(repeat_df[col].iloc[0]):
                                    repeat_df[col] = repeat_df[col].astype(str).str.replace('%', '').replace('', '0')
                                repeat_df[col] = pd.to_numeric(repeat_df[col], errors='coerce').fillna(0)
                        repeat_df.insert(0, 'Rank', range(1, len(repeat_df) + 1))
                        zip_file.writestr("Top_Products_Repeat_Rate.csv", repeat_df.to_csv(index=False))
                        
                        aov_df = pd.DataFrame(top_products_data['top_aov'])
                        for col in aov_df.columns:
                            if col not in ['Product Title', 'Merchant SKU'] and aov_df[col].dtype == 'object':
                                if 'AOV' in col or '$' in str(aov_df[col].iloc[0]):
                                    aov_df[col] = aov_df[col].astype(str).str.replace('$', '').replace('', '0')
                                aov_df[col] = pd.to_numeric(aov_df[col], errors='coerce').fillna(0)
                        aov_df.insert(0, 'Rank', range(1, len(aov_df) + 1))
                        zip_file.writestr("Top_Products_AOV.csv", aov_df.to_csv(index=False))
                        
                        ltv_df = pd.DataFrame(top_products_data['top_ltv'])
                        for col in ltv_df.columns:
                            if col not in ['Product Title', 'Merchant SKU'] and ltv_df[col].dtype == 'object':
                                if 'LTV' in col or '$' in str(ltv_df[col].iloc[0]):
                                    ltv_df[col] = ltv_df[col].astype(str).str.replace('$', '').replace('', '0')
                                ltv_df[col] = pd.to_numeric(ltv_df[col], errors='coerce').fillna(0)
                        ltv_df.insert(0, 'Rank', range(1, len(ltv_df) + 1))
                        zip_file.writestr("Top_Products_LTV.csv", ltv_df.to_csv(index=False))
                    except:
                        pass  # Skip if Top Products calculation fails
                
                zip_buffer.seek(0)
                
                st.download_button(
                    label="📥 Download ZIP File",
                    data=zip_buffer.getvalue(),
                    file_name="LTV_Analysis_Results.zip",
                    mime="application/zip"
                )
        else:
            st.error("Failed to process uploaded files. Please check the file format and try again.")
    
    else:
        st.info("👆 Please upload CSV files using the sidebar to get started.")
        
        # Show sample data format
        st.header("📋 Expected Data Format")
        st.markdown("""
        Your CSV files should contain the following columns:
        - **purchase_date**: Date of purchase
        - **buyer_email**: Customer email
        - **merchant_sku**: Product SKU
        - **amazon_order_id**: Order ID
        - **item_price**: Price of item
        - **shipped_quantity**: Quantity shipped
        - **currency**: Currency
        - **title**: Product title
        - And other relevant e-commerce data columns...
        """)

if __name__ == "__main__":
    main() 